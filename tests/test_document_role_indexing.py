"""Fail-closed document-role coverage for proposal-reference indexing."""

from pathlib import Path
from unittest.mock import Mock

import pytest
from docx import Document
from qdrant_client import QdrantClient

from proposal_ai_agent.embeddings.providers.base import EmbeddingProvider
from proposal_ai_agent.indexing import DocumentRole, IndexPipeline, IndexRequest
from proposal_ai_agent.indexing.exceptions import DocumentRoleAuthorizationError


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "documents" / "Reference Document" / "PRU_T72_Module_Breakdown.docx"
TEMPLATE_DIRECTORY = ROOT / "documents" / "Template Document"
TEMPLATE = next(TEMPLATE_DIRECTORY.rglob("*.docx"))


class CountingEmbeddingProvider(EmbeddingProvider):
    def __init__(self) -> None:
        self.calls = 0

    def embed(self, text: str):
        self.calls += 1
        return (0.1, 0.2, 0.3)

    def embed_batch(self, texts):
        values = tuple(texts)
        self.calls += 1
        return tuple((0.1, 0.2, 0.3) for _ in values)


def _document(path: Path) -> None:
    document = Document()
    document.add_heading("Reference", level=1)
    document.add_paragraph("Authorized reference knowledge.")
    document.save(path)


def _pipeline(client, provider: EmbeddingProvider) -> IndexPipeline:
    return IndexPipeline(client, provider, 3, "role-test-model")


def test_reference_role_is_accepted_and_propagated(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(IndexPipeline, "_REFERENCE_KNOWLEDGE_ROOT", tmp_path.resolve())
    source = tmp_path / "reference.docx"
    _document(source)
    client = QdrantClient(":memory:")
    provider = CountingEmbeddingProvider()

    result = _pipeline(client, provider).index(
        IndexRequest(source, "role_test", DocumentRole.REFERENCE_KNOWLEDGE)
    )

    assert result.statistics.documents_indexed == 1
    assert provider.calls == 1
    points, _ = client.scroll("role_test", limit=10, with_payload=True)
    assert points
    assert all(
        point.payload["document_role"] == DocumentRole.REFERENCE_KNOWLEDGE.value
        for point in points
    )


def test_declared_publishing_template_role_is_rejected_before_io(tmp_path: Path) -> None:
    source = tmp_path / "template.docx"
    _document(source)
    client = Mock()
    provider = CountingEmbeddingProvider()

    with pytest.raises(DocumentRoleAuthorizationError, match="PUBLISHING_TEMPLATE"):
        _pipeline(client, provider).index(
            IndexRequest(source, "role_test", DocumentRole.PUBLISHING_TEMPLATE)
        )

    assert provider.calls == 0
    client.assert_not_called()
    assert not client.method_calls


def test_direct_template_path_is_rejected_even_when_misdeclared_as_reference() -> None:
    client = Mock()
    provider = CountingEmbeddingProvider()

    with pytest.raises(DocumentRoleAuthorizationError) as error:
        _pipeline(client, provider).index(
            IndexRequest(TEMPLATE, "role_test", DocumentRole.REFERENCE_KNOWLEDGE)
        )

    message = str(error.value)
    assert str(TEMPLATE.resolve()) in message
    assert "authorized_root=" in message
    assert "expected_role=REFERENCE_KNOWLEDGE" in message
    assert provider.calls == 0
    assert not client.method_calls


def test_recursive_documents_input_fails_before_any_document_is_indexed() -> None:
    client = Mock()
    provider = CountingEmbeddingProvider()

    with pytest.raises(DocumentRoleAuthorizationError, match="authorized_root"):
        _pipeline(client, provider).index(
            IndexRequest(ROOT / "documents", "role_test", DocumentRole.REFERENCE_KNOWLEDGE)
        )

    assert REFERENCE.is_file()
    assert provider.calls == 0
    assert not client.method_calls


def test_missing_or_unknown_role_fails_closed(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="document_role is required"):
        IndexRequest(tmp_path, "role_test")
    with pytest.raises(ValueError, match="invalid document_role"):
        IndexRequest(tmp_path, "role_test", "UNKNOWN")  # type: ignore[arg-type]


def test_real_reference_path_remains_authorized() -> None:
    IndexPipeline._authorize_reference_paths(
        (REFERENCE,), DocumentRole.REFERENCE_KNOWLEDGE
    )


@pytest.mark.parametrize(
    "source",
    [
        TEMPLATE,
        ROOT / "data" / "raw" / "proposals" / "samples" / "Proposal Draft.docx",
    ],
)
def test_noncanonical_repository_sources_fail_before_embedding_or_qdrant(source: Path) -> None:
    assert source.is_file()
    client = Mock()
    provider = CountingEmbeddingProvider()

    with pytest.raises(DocumentRoleAuthorizationError, match="authorized_root"):
        _pipeline(client, provider).index(
            IndexRequest(source, "role_test", DocumentRole.REFERENCE_KNOWLEDGE)
        )

    assert provider.calls == 0
    assert not client.method_calls


def test_arbitrary_external_source_fails_before_embedding_or_qdrant(tmp_path: Path) -> None:
    source = tmp_path / "external.docx"
    _document(source)
    client = Mock()
    provider = CountingEmbeddingProvider()

    with pytest.raises(DocumentRoleAuthorizationError, match="authorized_root"):
        _pipeline(client, provider).index(
            IndexRequest(source, "role_test", DocumentRole.REFERENCE_KNOWLEDGE)
        )

    assert provider.calls == 0
    assert not client.method_calls


def test_resolved_symlink_escape_fails_closed(tmp_path: Path, monkeypatch) -> None:
    authorized_root = tmp_path / "Reference Document"
    authorized_root.mkdir()
    source = tmp_path / "external.docx"
    _document(source)
    link = authorized_root / "symlink-escape.docx"
    link.symlink_to(source)
    monkeypatch.setattr(IndexPipeline, "_REFERENCE_KNOWLEDGE_ROOT", authorized_root.resolve())
    client = Mock()
    provider = CountingEmbeddingProvider()

    with pytest.raises(DocumentRoleAuthorizationError, match="authorized_root"):
        _pipeline(client, provider).index(
            IndexRequest(link, "role_test", DocumentRole.REFERENCE_KNOWLEDGE)
        )

    assert provider.calls == 0
    assert not client.method_calls

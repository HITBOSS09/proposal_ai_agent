"""Tests for native PDF ingestion into the common document model."""

from pathlib import Path

import fitz
from docx import Document as DocxDocument

from proposal_ai_agent.ingestion.chunker import chunk_document
from proposal_ai_agent.ingestion.loader import (
    discover_document_files,
    load_docx_document,
    load_pdf_document,
)
from proposal_ai_agent.ingestion.models import Paragraph, Table
from proposal_ai_agent.ingestion.pdf_parser import PDFParserError, parse_pdf_document
from proposal_ai_agent.ingestion.validator import DocumentValidator


def _create_text_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "SECTION A: SUMMARY", fontsize=16)
    page.insert_text((72, 108), "This proposal describes the required delivery.", fontsize=10)
    document.save(path)
    document.close()


def _create_table_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page()
    for x_position in (72, 180, 288):
        page.draw_line((x_position, 72), (x_position, 132))
    for y_position in (72, 102, 132):
        page.draw_line((72, y_position), (288, y_position))
    page.insert_text((82, 92), "Name", fontsize=10)
    page.insert_text((190, 92), "Value", fontsize=10)
    page.insert_text((82, 122), "Scope", fontsize=10)
    page.insert_text((190, 122), "PDF", fontsize=10)
    document.save(path)
    document.close()


def _create_docx(path: Path) -> None:
    document = DocxDocument()
    document.add_paragraph("DOCX content")
    document.save(path)


def test_pdf_parser_creates_common_models_and_page_metadata(tmp_path):
    pdf_path = tmp_path / "proposal.pdf"
    _create_text_pdf(pdf_path)

    document = parse_pdf_document(pdf_path, source_document="proposal")

    assert document.title == "proposal"
    assert document.metadata.document_type == "proposal"
    assert len(document.sections) == 1
    assert document.sections[0].heading == "SECTION A: SUMMARY"
    assert len(document.sections[0].elements) == 1
    paragraph = document.sections[0].elements[0]
    assert isinstance(paragraph, Paragraph)
    assert paragraph.content == "This proposal describes the required delivery."
    assert paragraph.metadata.page_number == 1
    assert paragraph.metadata.source_file == "proposal.pdf"

    assert DocumentValidator().validate(document).is_valid
    chunks = chunk_document(document)
    assert len(chunks) == 1
    assert "This proposal describes the required delivery." in chunks[0].text


def test_pdf_parser_detects_tables(tmp_path):
    pdf_path = tmp_path / "table.pdf"
    _create_table_pdf(pdf_path)

    document = parse_pdf_document(pdf_path)

    assert len(document.elements) == 1
    table = document.elements[0]
    assert isinstance(table, Table)
    assert [[cell.content for cell in row.cells] for row in table.rows] == [
        ["Name", "Value"],
        ["Scope", "PDF"],
    ]
    assert table.metadata.page_number == 1


def test_pdf_parser_rejects_textless_pdf(tmp_path):
    pdf_path = tmp_path / "empty.pdf"
    document = fitz.open()
    document.new_page()
    document.save(pdf_path)
    document.close()

    try:
        parse_pdf_document(pdf_path)
    except PDFParserError as error:
        assert "no extractable text" in str(error)
    else:
        raise AssertionError("Expected a textless PDF to be rejected")


def test_loader_discovers_and_loads_pdf_files(tmp_path):
    pdf_path = tmp_path / "proposal.pdf"
    docx_path = tmp_path / "proposal.docx"
    _create_text_pdf(pdf_path)
    _create_docx(docx_path)
    (tmp_path / "ignored.txt").write_text("not a document")

    assert discover_document_files(tmp_path) == [docx_path, pdf_path]

    docx_document = load_docx_document(docx_path)
    assert docx_document.paragraphs[0].text == "DOCX content"

    pdf_document = load_pdf_document(pdf_path)
    try:
        assert pdf_document.page_count == 1
    finally:
        pdf_document.close()

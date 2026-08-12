"""
Tests for the Document Parser.

Comprehensive test suite covering:
- Heading detection and hierarchy
- Paragraph and table parsing
- Section nesting
- Order index preservation
- Metadata propagation
- Edge cases (empty paragraphs, flat documents, etc.)
"""

import pytest
from pathlib import Path
from uuid import UUID
from docx import Document as DocxDocument
from docx.shared import Pt

from proposal_ai_agent.ingestion.parser import (
    DocumentParser,
    extract_heading_level,
    is_heading_paragraph,
    parse_paragraph,
    parse_table,
)
from proposal_ai_agent.ingestion.models import (
    Document,
    Section,
    Paragraph,
    Table,
)


# ============================================================================
# FIXTURES
# ============================================================================

@pytest.fixture
def temp_docx_file(tmp_path):
    """Create a temporary DOCX file for testing."""
    docx_path = tmp_path / "test_document.docx"
    return docx_path


@pytest.fixture
def parser():
    """Create a DocumentParser instance."""
    return DocumentParser()


# ============================================================================
# HEADING DETECTION TESTS
# ============================================================================

class TestHeadingDetection:
    """Tests for heading level extraction and detection."""
    
    def test_extract_heading_level_heading1(self):
        """Test extraction of Heading 1 style."""
        level = extract_heading_level("Heading 1")
        assert level == 1
    
    def test_extract_heading_level_heading2(self):
        """Test extraction of Heading 2 style."""
        level = extract_heading_level("Heading 2")
        assert level == 2
    
    def test_extract_heading_level_heading3(self):
        """Test extraction of Heading 3 style."""
        level = extract_heading_level("Heading 3")
        assert level == 3
    
    def test_extract_heading_level_case_insensitive(self):
        """Test that heading extraction is case-insensitive."""
        level = extract_heading_level("heading 1")
        assert level == 1
    
    def test_extract_heading_level_with_whitespace(self):
        """Test heading extraction with extra whitespace."""
        level = extract_heading_level("  Heading 2  ")
        assert level == 2
    
    def test_extract_heading_level_non_heading(self):
        """Test that non-heading styles return None."""
        level = extract_heading_level("Normal")
        assert level is None
    
    def test_extract_heading_level_none_style(self):
        """Test that None returns None."""
        level = extract_heading_level(None)
        assert level is None
    
    def test_extract_heading_level_invalid_number(self):
        """Test that invalid heading numbers return None."""
        level = extract_heading_level("Heading 99")
        assert level is None
    
    def test_extract_heading_level_malformed(self):
        """Test that malformed heading strings return None."""
        level = extract_heading_level("Heading ABC")
        assert level is None


# ============================================================================
# DOCUMENT PARSING TESTS
# ============================================================================

class TestDocumentParsing:
    """Tests for complete document parsing."""
    
    def test_parse_simple_flat_document(self, parser, temp_docx_file):
        """Test parsing a document without headings."""
        # Create simple document
        docx = DocxDocument()
        docx.add_paragraph("First paragraph.", style="Normal")
        docx.add_paragraph("Second paragraph.", style="Normal")
        docx.save(temp_docx_file)
        
        # Reload and parse
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file)
        
        assert parsed.title == "Untitled Document"
        assert len(parsed.sections) == 0
        assert len(parsed.elements) == 2
        
        # Check that paragraphs have correct order_index
        assert parsed.elements[0].metadata.order_index == 0
        assert parsed.elements[1].metadata.order_index == 1
    
    def test_parse_document_with_heading(self, parser, temp_docx_file):
        """Test parsing document with a single heading and content."""
        docx = DocxDocument()
        docx.add_paragraph("Introduction", style="Heading 1")
        docx.add_paragraph("This is the introduction.", style="Normal")
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file)
        
        assert len(parsed.sections) == 1
        assert parsed.sections[0].heading == "Introduction"
        assert parsed.sections[0].section_level == 1
        assert len(parsed.sections[0].elements) == 1
    
    def test_parse_hierarchical_sections(self, parser, temp_docx_file):
        """Test parsing document with nested sections (H1 -> H2 -> H3)."""
        docx = DocxDocument()
        docx.add_paragraph("Main Section", style="Heading 1")
        docx.add_paragraph("Content in main.", style="Normal")
        docx.add_paragraph("Subsection", style="Heading 2")
        docx.add_paragraph("Content in subsection.", style="Normal")
        docx.add_paragraph("Sub-subsection", style="Heading 3")
        docx.add_paragraph("Content in sub-subsection.", style="Normal")
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file)
        
        # Check top-level section
        assert len(parsed.sections) == 1
        main = parsed.sections[0]
        assert main.heading == "Main Section"
        assert len(main.elements) == 1
        
        # Check subsection
        assert len(main.subsections) == 1
        sub = main.subsections[0]
        assert sub.heading == "Subsection"
        assert sub.section_level == 2
        assert len(sub.elements) == 1
        
        # Check sub-subsection
        assert len(sub.subsections) == 1
        subsub = sub.subsections[0]
        assert subsub.heading == "Sub-subsection"
        assert subsub.section_level == 3
    
    def test_parse_multiple_top_level_sections(self, parser, temp_docx_file):
        """Test parsing document with multiple H1 sections."""
        docx = DocxDocument()
        docx.add_paragraph("Section 1", style="Heading 1")
        docx.add_paragraph("Content in section 1.", style="Normal")
        docx.add_paragraph("Section 2", style="Heading 1")
        docx.add_paragraph("Content in section 2.", style="Normal")
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file)
        
        assert len(parsed.sections) == 2
        assert parsed.sections[0].heading == "Section 1"
        assert parsed.sections[1].heading == "Section 2"
    
    def test_parse_skip_empty_paragraphs(self, parser, temp_docx_file):
        """Test that empty paragraphs are skipped."""
        docx = DocxDocument()
        docx.add_paragraph("First.", style="Normal")
        docx.add_paragraph("", style="Normal")  # Empty
        docx.add_paragraph("   ", style="Normal")  # Whitespace only
        docx.add_paragraph("Third.", style="Normal")
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file)
        
        # Should only have 2 paragraphs (empty ones skipped)
        assert len(parsed.elements) == 2
        assert parsed.elements[0].content == "First."
        assert parsed.elements[1].content == "Third."
    
    def test_parse_preserve_order_index(self, parser, temp_docx_file):
        """Test that order_index is preserved across mixed content."""
        docx = DocxDocument()
        docx.add_paragraph("Heading", style="Heading 1")
        docx.add_paragraph("Para 1", style="Normal")
        table = docx.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Cell 1"
        table.rows[0].cells[1].text = "Cell 2"
        docx.add_paragraph("Para 2", style="Normal")
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file)
        
        section = parsed.sections[0]
        # Order should be: Para 1 (0), Table (1), Para 2 (2)
        assert section.elements[0].metadata.order_index == 0
        assert section.elements[1].metadata.order_index == 1
        assert section.elements[2].metadata.order_index == 2
    
    def test_parse_metadata_propagation(self, parser, temp_docx_file):
        """Test that metadata is correctly propagated to elements."""
        docx = DocxDocument()
        docx.add_paragraph("Heading", style="Heading 1")
        docx.add_paragraph("Content", style="Normal")
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(
            docx,
            temp_docx_file,
            source_document="test_doc",
            author="Test Author"
        )
        
        # Check document metadata
        assert parsed.metadata.author == "Test Author"
        assert parsed.metadata.document_type == "test_doc"
        assert str(parsed.metadata.source) == str(temp_docx_file)
        
        # Check element metadata
        element = parsed.sections[0].elements[0]
        assert element.metadata.document_id == parsed.document_id
        assert element.metadata.source_file == temp_docx_file.name
        assert element.metadata.source_document == "test_doc"
    
    def test_parse_document_title_override(self, parser, temp_docx_file):
        """Test that document title can be overridden."""
        docx = DocxDocument()
        docx.add_paragraph("Auto Title", style="Heading 1")
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file, document_title="Custom Title")
        
        assert parsed.title == "Custom Title"
    
    def test_parse_document_title_from_first_heading(self, parser, temp_docx_file):
        """Test that document title is inferred from first heading."""
        docx = DocxDocument()
        docx.add_paragraph("Auto Title", style="Heading 1")
        docx.add_paragraph("Content", style="Normal")
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file)
        
        assert parsed.title == "Auto Title"
    
    def test_parse_table_structure(self, parser, temp_docx_file):
        """Test that tables are correctly parsed with structure."""
        docx = DocxDocument()
        docx.add_paragraph("Tables Section", style="Heading 1")
        table = docx.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Header 2"
        table.rows[1].cells[0].text = "Data 1"
        table.rows[1].cells[1].text = "Data 2"
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file)
        
        section = parsed.sections[0]
        table_elem = section.elements[0]
        
        assert isinstance(table_elem, Table)
        assert len(table_elem.rows) == 3
        assert len(table_elem.rows[0].cells) == 2
        assert table_elem.rows[0].cells[0].content == "Header 1"
        assert table_elem.rows[0].cells[0].cell_type == "header"
        assert table_elem.rows[1].cells[0].cell_type == "data"
    
    def test_parse_root_and_sectioned_content(self, parser, temp_docx_file):
        """Test document with both root-level and sectioned content."""
        docx = DocxDocument()
        docx.add_paragraph("Preamble", style="Normal")
        docx.add_paragraph("Main Section", style="Heading 1")
        docx.add_paragraph("Section content", style="Normal")
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file)
        
        # Check root element
        assert len(parsed.elements) == 1
        assert parsed.elements[0].content == "Preamble"
        
        # Check section
        assert len(parsed.sections) == 1
        assert parsed.sections[0].heading == "Main Section"


# ============================================================================
# INTEGRATION TESTS
# ============================================================================

class TestParserIntegration:
    """Integration tests with real document scenarios."""
    
    def test_parse_complex_document(self, parser, temp_docx_file):
        """Test parsing a complex document with mixed content and hierarchy."""
        docx = DocxDocument()
        
        # Introduction (root level)
        docx.add_paragraph("This is the introduction to the proposal.", style="Normal")
        
        # Section 1
        docx.add_paragraph("Executive Summary", style="Heading 1")
        docx.add_paragraph("Summary content here.", style="Normal")
        
        # Section 1.1
        docx.add_paragraph("Key Metrics", style="Heading 2")
        table = docx.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Value"
        table.rows[1].cells[0].text = "Revenue"
        table.rows[1].cells[1].text = "$1M"
        
        # Section 1.1.1
        docx.add_paragraph("Analysis", style="Heading 3")
        docx.add_paragraph("Detailed analysis of metrics.", style="Normal")
        
        # Section 2
        docx.add_paragraph("Implementation Plan", style="Heading 1")
        docx.add_paragraph("Steps to implement:", style="Normal")
        
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file)
        
        # Verify structure
        assert len(parsed.elements) == 1  # Preamble
        assert len(parsed.sections) == 2  # Two H1 sections
        
        # Check first section hierarchy
        section1 = parsed.sections[0]
        assert section1.heading == "Executive Summary"
        assert len(section1.elements) == 1
        assert len(section1.subsections) == 1
        
        # Check subsection
        subsection = section1.subsections[0]
        assert subsection.heading == "Key Metrics"
        assert len(subsection.elements) == 1
        assert isinstance(subsection.elements[0], Table)
        assert len(subsection.subsections) == 1
        
        # Check sub-subsection
        subsubsection = subsection.subsections[0]
        assert subsubsection.heading == "Analysis"
        assert subsubsection.section_level == 3
    
    def test_parse_preserves_text_unchanged(self, parser, temp_docx_file):
        """Test that text content is NOT modified during parsing."""
        test_content = "  This has   multiple   spaces  \nAnd newlines\t\tand\ttabs  "
        docx = DocxDocument()
        docx.add_paragraph(test_content, style="Normal")
        docx.save(temp_docx_file)
        
        docx = DocxDocument(temp_docx_file)
        parsed = parser.parse(docx, temp_docx_file)
        
        # Text should be unchanged
        assert parsed.elements[0].content == test_content


# ============================================================================
# RUN TESTS
# ============================================================================

if __name__ == "__main__":
    pytest.main([__file__, "-v"])

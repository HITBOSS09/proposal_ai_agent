"""Phase 4C.1.5 crisp-output visual and package fidelity contracts."""

import inspect
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import pytest

from proposal_ai_agent.proposal_generation.publishing import (
    BindingAuthorizationError,
    CertifiedContracts,
    CertifiedTemplateBinding,
    file_sha256,
)
from scripts.run_phase4c_visual_fidelity_hardening import (
    CONTRACTS,
    EXPECTED_SHA,
    OUTPUT,
    PACKAGE_DIFF,
    SOURCE,
    package_diff,
    run_hardening,
)


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def _c14n(nodes) -> tuple[bytes, ...]:
    return tuple(etree.tostring(node, method="c14n", exclusive=True) for node in nodes)


def _properties(node) -> tuple[bytes, ...]:
    return _c14n(node.xpath(".//w:pPr | .//w:rPr | .//w:tblPr | .//w:tblGrid | .//w:trPr | .//w:tcPr"))


def test_crisp_generation_preserves_source_and_reopens(tmp_path: Path) -> None:
    before = file_sha256(SOURCE)
    output = run_hardening(tmp_path / "crisp.docx")
    reopened = Document(output)
    assert before == file_sha256(SOURCE) == EXPECTED_SHA
    assert len(reopened.sections) == 17
    assert len(reopened.tables) == 20


def test_existing_certified_text_table_clone_and_duplicate_isolation_survive() -> None:
    document = Document(OUTPUT)
    text = "".join(document.element.body.xpath(".//w:t/text()"))
    assert "SKYSHIELD-X ACCEPTANCE PRU" in text
    assert "Synthetic Sensor Unit" in document.tables[1].rows[0].cells[1].text
    assert "Cloned Sensor Node" in document.tables[13].rows[6].cells[1].text
    assert "Isolation Beacon" in document.tables[15].rows[1].cells[1].text
    assert "Isolation Beacon" not in " ".join(cell.text for row in document.tables[13].rows for cell in row.cells)


def test_text_replacement_preserves_paragraph_and_representative_run_properties() -> None:
    source = Document(SOURCE)
    output = Document(OUTPUT)
    source_nodes = tuple(source.element.body.iterchildren())
    output_nodes = tuple(output.element.body.iterchildren())
    for index in (4, 13, 127):
        assert _c14n(output_nodes[index].xpath(".//w:pPr")) == _c14n(source_nodes[index].xpath(".//w:pPr"))
        assert _c14n(output_nodes[index].xpath(".//w:rPr")) == _c14n(source_nodes[index].xpath(".//w:rPr"))


def test_table_grid_row_cell_alignment_and_merge_properties_survive() -> None:
    source = Document(SOURCE)
    output = Document(OUTPUT)
    for table_index in range(20):
        assert etree.tostring(output.tables[table_index]._tbl.tblPr, method="c14n") == etree.tostring(source.tables[table_index]._tbl.tblPr, method="c14n")
        assert etree.tostring(output.tables[table_index]._tbl.tblGrid, method="c14n") == etree.tostring(source.tables[table_index]._tbl.tblGrid, method="c14n")
    for table_index, row_index in ((1, 0), (12, 1), (13, 1), (15, 1)):
        source_row = source.tables[table_index].rows[row_index]._tr
        output_row = output.tables[table_index].rows[row_index]._tr
        assert _properties(output_row) == _properties(source_row)
        assert _c14n(output_row.xpath(".//w:vAlign | .//w:jc")) == _c14n(source_row.xpath(".//w:vAlign | .//w:jc"))
    assert _c14n(output.element.body.xpath(".//w:gridSpan")) == _c14n(source.element.body.xpath(".//w:gridSpan"))
    assert _properties(output.tables[13].rows[6]._tr) == _properties(source.tables[13].rows[2]._tr)


def test_sections_headers_footers_fields_and_relationship_references_survive() -> None:
    with ZipFile(SOURCE) as source, ZipFile(OUTPUT) as output:
        source_doc = etree.fromstring(source.read("word/document.xml"))
        output_doc = etree.fromstring(output.read("word/document.xml"))
        assert _c14n(output_doc.xpath(".//w:sectPr", namespaces=NS)) == _c14n(source_doc.xpath(".//w:sectPr", namespaces=NS))
        for index in range(1, 6):
            source_footer = etree.fromstring(source.read(f"word/footer{index}.xml"))
            output_footer = etree.fromstring(output.read(f"word/footer{index}.xml"))
            assert etree.tostring(output_footer, method="c14n") == etree.tostring(source_footer, method="c14n")
            assert output_footer.xpath(".//w:instrText/text()", namespaces=NS) == source_footer.xpath(".//w:instrText/text()", namespaces=NS)


def test_branding_survives_and_all_certified_project_diagrams_become_exact_placeholders() -> None:
    contracts = CertifiedContracts.load(CONTRACTS)
    document = Document(OUTPUT)
    nodes = tuple(document.element.body.iterchildren())
    with ZipFile(SOURCE) as source, ZipFile(OUTPUT) as output:
        assert output.read("word/media/image1.png") == source.read("word/media/image1.png")
        assert not any(name.endswith(".jpeg") for name in output.namelist())
    for location in contracts.schema["manual_diagram_locations"]:
        node = nodes[location["locator"]["body_node_index"]]
        assert "".join(node.xpath(".//w:t/text()")) == "[ INSERT DIAGRAM ]"
        assert not node.xpath(".//w:drawing")


def test_uncertified_and_branding_images_are_rejected_by_diagram_api(tmp_path: Path) -> None:
    contracts = CertifiedContracts.load(CONTRACTS)
    binding = CertifiedTemplateBinding(SOURCE, contracts, tmp_path / "copy.docx")
    for location_id in ("cover_logo", "image_7", "unknown_visual"):
        with pytest.raises(BindingAuthorizationError):
            binding.replace_certified_manual_diagram(location_id)


def test_no_dynamic_or_generated_diagram_or_generic_reconstruction_api_exists() -> None:
    source = inspect.getsource(CertifiedTemplateBinding)
    assert "generate_diagram" not in source
    assert "image_generation" not in source
    assert "add_picture" not in source
    assert "add_table" not in source
    assert not hasattr(CertifiedTemplateBinding, "replace_image")


def test_no_markdown_and_no_section_normalization() -> None:
    document = Document(OUTPUT)
    text = "".join(document.element.body.xpath(".//w:t/text()"))
    assert "#" not in text and "**" not in text
    assert len(document.sections) == 17
    assert all(not section["normalization_permitted"] for section in CertifiedContracts.load(CONTRACTS).policy["word_sections"])


def test_package_diff_contains_only_explained_certified_changes() -> None:
    diff = package_diff(SOURCE, OUTPUT)
    assert diff["unexplained_changed_parts"] == []
    assert diff["unexplained_removed_parts"] == []
    assert diff["added_parts"] == []
    assert diff["changed_parts"] == [
        "word/_rels/document.xml.rels", "word/document.xml", "word/header1.xml",
        "word/header2.xml", "word/header3.xml", "word/header4.xml", "word/header5.xml",
    ]
    assert PACKAGE_DIFF.is_file()

"""Fidelity and safety tests for low-level template-preserving DOCX edits."""

from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from lxml import etree
import pytest

from proposal_ai_agent.proposal_generation.publishing import (
    OpenXmlEditor,
    ReusableComponent,
    SemanticField,
    UnsafeOpenXmlOperation,
    assert_footer_contract,
    file_sha256,
    pru_template_semantic_map,
)


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "documents/PRU_T72_Module_Breakdown.docx"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def _package_counts(path: Path) -> dict[str, int]:
    with ZipFile(path) as package:
        names = package.namelist()
        document = etree.fromstring(package.read("word/document.xml"))
        header_parts = [name for name in names if name.startswith("word/header") and name.endswith(".xml")]
        footer_parts = [name for name in names if name.startswith("word/footer") and name.endswith(".xml")]
        header_footer_drawings = sum(
            len(etree.fromstring(package.read(name)).xpath(".//w:drawing", namespaces=NS))
            for name in header_parts + footer_parts
        )
        return {
            "sections": len(document.xpath(".//w:sectPr", namespaces=NS)),
            "tables": len(document.xpath("./w:body/w:tbl", namespaces=NS)),
            "headers": len(header_parts),
            "footers": len(footer_parts),
            "media": len([name for name in names if name.startswith("word/media/")]),
            "body_drawings": len(document.xpath(".//w:body//w:drawing", namespaces=NS)),
            "header_footer_drawings": header_footer_drawings,
        }


def _banner_properties(node) -> tuple[tuple[str, ...], tuple[str, ...]]:
    fills = tuple(node.xpath('.//a:solidFill/a:srgbClr/@val'))
    outlines = tuple(node.xpath('.//a:ln/a:solidFill/a:srgbClr/@val'))
    return fills, outlines


def _named_body_drawing(path: Path, name: str) -> bytes:
    with ZipFile(path) as package:
        document = etree.fromstring(package.read("word/document.xml"))
    matches = document.xpath(
        f'.//w:body//w:drawing[.//wp:docPr[@name="{name}"]]',
        namespaces=NS | {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"},
    )
    assert len(matches) == 1
    return etree.tostring(matches[0], method="c14n")


def test_run_aware_replacement_preserves_run_and_paragraph_properties(tmp_path: Path) -> None:
    fixture = Document()
    paragraph = fixture.add_paragraph()
    first = paragraph.add_run("Formatted ")
    first.bold = True
    first.font.color.rgb = RGBColor(0xE9, 0x6F, 0x30)
    second = paragraph.add_run("project title")
    second.italic = True
    source_rpr = [etree.tostring(run._r.rPr) for run in paragraph.runs]
    source_ppr = None if paragraph._p.pPr is None else etree.tostring(paragraph._p.pPr)

    OpenXmlEditor._replace_in_container(paragraph._p, "Formatted project title", "New proposal heading")

    assert paragraph.text == "New proposal heading"
    assert [etree.tostring(run._r.rPr) for run in paragraph.runs] == source_rpr
    assert (None if paragraph._p.pPr is None else etree.tostring(paragraph._p.pPr)) == source_ppr


def test_real_template_safe_edits_preserve_package_and_visual_contracts(tmp_path: Path) -> None:
    source_hash = file_sha256(TEMPLATE)
    before = _package_counts(TEMPLATE)
    untouched_logo = _named_body_drawing(TEMPLATE, "Image 4")
    untouched_diagram = _named_body_drawing(TEMPLATE, "Image 7")
    semantic_map = pru_template_semantic_map(TEMPLATE)
    editor = OpenXmlEditor(TEMPLATE, tmp_path / "working.docx", semantic_map)

    metadata_table = editor.document.tables[0]
    table_properties = etree.tostring(metadata_table._tbl.tblPr)
    value_cell = metadata_table.rows[0].cells[1]
    cell_properties = etree.tostring(value_cell._tc.tcPr)
    editor.replace_table_cell_text(
        semantic_map.target(SemanticField.DOCUMENT_NUMBER).locators[0], "BDIL-TEST-001"
    )
    assert etree.tostring(metadata_table._tbl.tblPr) == table_properties
    assert etree.tostring(value_cell._tc.tcPr) == cell_properties
    assert editor.replace_text(
        semantic_map.target(SemanticField.DOCUMENT_TITLE).locators[0],
        "TECHNICAL PROJECT PROPOSAL",
    ) == 2

    orange_bom = editor.document.tables[13]
    prototype_row = orange_bom.rows[1]._tr
    prototype_trpr = etree.tostring(prototype_row.trPr)
    prototype_tcpr = tuple(etree.tostring(cell.tcPr) for cell in prototype_row.tc_lst)
    cloned_row = editor.clone_table_row(13, 1)
    assert etree.tostring(cloned_row.trPr) == prototype_trpr
    assert tuple(etree.tostring(cell.tcPr) for cell in cloned_row.tc_lst) == prototype_tcpr

    banner_locator = semantic_map.prototype(ReusableComponent.MODULE_BANNER).locator
    banner_node = editor.resolve_body_node(banner_locator)
    banner_properties = _banner_properties(banner_node)
    banner_clone = editor.clone_banner(banner_locator)
    assert _banner_properties(banner_clone) == banner_properties
    editor.insert_body_node(banner_node, banner_clone, "after")
    editor.replace_drawing_text(banner_locator, "MODULE 8 | SYNTHETIC TEST SYSTEM")
    assert _banner_properties(banner_node) == banner_properties
    assert "E96F30" in banner_properties[0]
    assert "C2541B" in banner_properties[1]

    for locator in semantic_map.target(SemanticField.HEADER_PROJECT_TEXT).locators:
        assert editor.replace_drawing_text(locator, "TEST PROGRAMME — TECHNICAL PROPOSAL") == 2
    editor.assert_footers_preserved("Bharati Defence and Infrastructure Limited")

    removable = deepcopy(editor.body_nodes[13])
    reference = editor.body_nodes[13]
    editor.insert_body_node(reference, removable, "before")
    editor.remove_body_node(removable)
    section_node = next(node for node in editor.body_nodes if node.xpath("./w:pPr/w:sectPr"))
    with pytest.raises(UnsafeOpenXmlOperation):
        editor.remove_body_node(section_node)

    output = editor.save(tmp_path / "edited.docx")
    reopened = Document(output)
    assert_footer_contract(reopened, "Bharati Defence and Infrastructure Limited")
    after = _package_counts(output)

    assert reopened.tables[0].rows[0].cells[1].text == "BDIL-TEST-001"
    assert "TECHNICAL PROJECT PROPOSAL" in "".join(reopened.element.body.xpath(".//w:t/text()"))
    assert len(reopened.tables[13].rows) == len(orange_bom.rows)
    assert after == before | {"body_drawings": before["body_drawings"] + 1}
    assert _named_body_drawing(output, "Image 4") == untouched_logo
    assert _named_body_drawing(output, "Image 7") == untouched_diagram
    assert file_sha256(TEMPLATE) == source_hash

    body = reopened.element.body
    docpr_ids = body.xpath(".//wp:docPr/@id")
    vml_ids = [
        element.get("id") for element in body.iter("{urn:schemas-microsoft-com:vml}shape")
    ]
    assert len(docpr_ids) == len(set(docpr_ids))
    assert len(vml_ids) == len(set(vml_ids))

    header_parts = [
        part for part in reopened.part.package.parts
        if str(part.partname).startswith("/word/header")
    ]
    assert len(header_parts) == 5
    for part in header_parts:
        text = "".join(part.element.xpath(".//w:t/text()"))
        assert text.count("TEST PROGRAMME — TECHNICAL PROPOSAL") == 2
        assert len(part.element.xpath(".//w:drawing")) == 1


def test_source_template_cannot_be_used_as_working_or_output_path(tmp_path: Path) -> None:
    semantic_map = pru_template_semantic_map(TEMPLATE)
    with pytest.raises(ValueError):
        OpenXmlEditor(TEMPLATE, TEMPLATE, semantic_map)

    editor = OpenXmlEditor(TEMPLATE, tmp_path / "working.docx", semantic_map)
    with pytest.raises(UnsafeOpenXmlOperation):
        editor.save(TEMPLATE)

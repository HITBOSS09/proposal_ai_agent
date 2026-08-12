"""End-to-end indexing-runtime integration tests using in-memory Qdrant."""

from pathlib import Path

import pytest
from docx import Document
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, PayloadSchemaType, VectorParams

from proposal_ai_agent.embeddings.providers import MockEmbeddingProvider
from proposal_ai_agent.indexing import CollectionManager, DocumentRole, IndexPipeline, IndexRequest
from proposal_ai_agent.indexing.exceptions import CollectionConfigurationMismatch


def _proposal(path: Path, heading: str, paragraph: str) -> None:
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(paragraph)
    document.save(path)


def _pipeline(client: QdrantClient) -> IndexPipeline:
    return IndexPipeline(client, MockEmbeddingProvider(dimensions=3), 3, "test-embedding-model")


def test_collection_manager_creates_cosine_collection_with_runtime_metadata() -> None:
    client = QdrantClient(":memory:")
    manager = CollectionManager(client, "bdil_demo", 3, "test-embedding-model")

    assert manager.ensure_ready() is True
    collection = client.get_collection("bdil_demo")
    assert collection.config.params.vectors.size == 3
    assert collection.config.params.vectors.distance is Distance.COSINE
    assert collection.config.metadata["embedding_model"] == "test-embedding-model"
    assert set(manager._PAYLOAD_INDEXES) == {
        "document_id", "document_title", "section_path", "page_number", "chunk_id",
        "source_document", "version", "document_fingerprint", "document_role",
    }
    assert manager._PAYLOAD_INDEXES["document_role"] is PayloadSchemaType.KEYWORD


def test_pipeline_indexes_batches_and_reuses_its_collection(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(IndexPipeline, "_REFERENCE_KNOWLEDGE_ROOT", tmp_path.resolve())
    _proposal(tmp_path / "first.docx", "Security", "Cybersecurity controls are mandatory.")
    _proposal(tmp_path / "second.docx", "Operations", "Service continuity requirements apply.")
    client = QdrantClient(":memory:")

    result = _pipeline(client).index(
        IndexRequest(tmp_path, "bdil_demo", DocumentRole.REFERENCE_KNOWLEDGE, batch_size=1)
    )

    assert result.collection_created is True
    assert result.statistics.documents_indexed == 2
    assert result.statistics.chunks_indexed == 2
    assert result.statistics.embeddings_generated == 2
    assert result.statistics.vectors_uploaded == 2
    assert result.statistics.documents_failed == 0
    assert client.count("bdil_demo", exact=True).count == 2
    payloads, _ = client.scroll("bdil_demo", limit=10, with_payload=True, with_vectors=False)
    assert {payload.payload["document_title"] for payload in payloads} == {"first", "second"}
    assert all("document_fingerprint" in payload.payload for payload in payloads)
    assert all(
        payload.payload["document_role"] == DocumentRole.REFERENCE_KNOWLEDGE.value
        for payload in payloads
    )

    reused = _pipeline(client).index(
        IndexRequest(tmp_path, "bdil_demo", DocumentRole.REFERENCE_KNOWLEDGE, batch_size=1)
    )
    assert reused.collection_created is False
    assert reused.statistics.documents_indexed == 0
    assert reused.statistics.documents_skipped == 2
    assert client.count("bdil_demo", exact=True).count == 2


def test_collection_manager_rejects_dimension_or_model_mismatch() -> None:
    client = QdrantClient(":memory:")
    CollectionManager(client, "bdil_demo", 3, "model-a").ensure_ready()

    with pytest.raises(CollectionConfigurationMismatch, match="embedding model"):
        CollectionManager(client, "bdil_demo", 3, "model-b").ensure_ready()
    with pytest.raises(CollectionConfigurationMismatch, match="dimension 4"):
        CollectionManager(client, "bdil_demo", 4, "model-a").ensure_ready()

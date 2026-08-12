"""Tests for template-owned DOCX compilation from semantic composition."""

from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
import pytest

from proposal_ai_agent.proposal_generation.composition import CompositionEngine
from proposal_ai_agent.proposal_generation.document_plan import ProposalPlan, SectionPlan, SectionRole
from proposal_ai_agent.proposal_generation.docx_compiler import DOCXCompiler, MissingTemplateStyle
from proposal_ai_agent.proposal_generation.proposal_ir import (
    BulletList,
    Callout,
    Heading,
    KnowledgeReference,
    Paragraph,
    ProposalDocument,
    RequirementMatrix,
    RequirementMatrixEntry,
    Section,
    Table,
    VisualPlaceholder,
)
from proposal_ai_agent.proposal_generation.word_style_contract import (
    BodyStyle,
    BulletStyle,
    CalloutStyle,
    CaptionStyle,
    CoverTitleStyle,
    FooterStyle,
    HeaderStyle,
    Heading1Style,
    Heading2Style,
    Heading3Style,
    ModuleBannerStyle,
    PageNumberStyle,
    ProposalTitleStyle,
    RequirementMatrixStyle,
    TableCellStyle,
    TableHeaderStyle,
)


def _styles():
    return {
        "proposal_title": ProposalTitleStyle(),
        "cover_title": CoverTitleStyle(),
        "module_banner": ModuleBannerStyle(),
        "heading_1": Heading1Style(),
        "heading_2": Heading2Style(),
        "heading_3": Heading3Style(),
        "body": BodyStyle(),
        "bullet": BulletStyle(),
        "table_header": TableHeaderStyle(),
        "table_cell": TableCellStyle(),
        "requirement_matrix": RequirementMatrixStyle(),
        "callout": CalloutStyle(),
        "caption": CaptionStyle(),
        "header": HeaderStyle(),
        "footer": FooterStyle(),
        "page_number": PageNumberStyle(),
    }


def _template(path: Path) -> None:
    document = DocxDocument()
    for style in _styles().values():
        document.styles.add_style(style.template_style_name, WD_STYLE_TYPE.PARAGRAPH)
    document.sections[0].header.paragraphs[0].text = "Template header"
    document.sections[0].footer.paragraphs[0].text = "Template footer"
    document.save(path)


def _composition():
    proposal = ProposalDocument(
        proposal_id="proposal-001",
        title="Autonomous Perimeter Monitoring Proposal",
        sections=(
            Section(section_id="cover", heading=Heading(text="Cover", level=1)),
            Section(section_id="contents", heading=Heading(text="Contents", level=1)),
            Section(
                section_id="module-solution",
                heading=Heading(text="Solution Module", level=1),
                children=(
                    Section(
                        section_id="solution",
                        heading=Heading(text="Solution Overview", level=2),
                        blocks=(
                            Paragraph(text="Persistent monitoring.", reference_ids=("REF-1",)),
                            BulletList(items=("EO/IR observation",)),
                            Table(headers=("Requirement", "Response"), rows=(("R-1", "Covered"),)),
                            VisualPlaceholder(visual_id="VIS-1", description="System architecture", caption="Platform overview"),
                            Callout(label="Operational note", text="Site survey required."),
                            RequirementMatrix(entries=(RequirementMatrixEntry(
                                requirement_id="R-1", requirement="EO/IR", response="Included", evidence_reference_ids=("REF-1",)
                            ),)),
                        ),
                    ),
                ),
            ),
            Section(section_id="appendix-a", heading=Heading(text="Appendix A", level=1)),
            Section(section_id="references", heading=Heading(text="References", level=1)),
        ),
        references=(KnowledgeReference(reference_id="REF-1", title="Payload specification", source="Technical library"),),
    )
    plan = ProposalPlan(
        proposal_id="proposal-001",
        sections=(
            SectionPlan(section_id="cover", role=SectionRole.COVER, include_in_toc=False, numbering_enabled=False),
            SectionPlan(section_id="contents", role=SectionRole.TABLE_OF_CONTENTS, numbering_enabled=False),
            SectionPlan(section_id="module-solution", role=SectionRole.MODULE),
            SectionPlan(section_id="solution", role=SectionRole.BODY),
            SectionPlan(section_id="appendix-a", role=SectionRole.APPENDIX),
            SectionPlan(section_id="references", role=SectionRole.REFERENCES),
        ),
    )
    return CompositionEngine().compose(plan, proposal)


def test_docx_compiler_generates_openable_template_styled_document(tmp_path: Path) -> None:
    template = tmp_path / "master_template.docx"
    output = tmp_path / "proposal.docx"
    _template(template)

    result = DOCXCompiler(_styles()).compile(_composition(), template, output)
    document = DocxDocument(result)
    paragraphs = {paragraph.text: paragraph.style.name for paragraph in document.paragraphs}

    assert result == output
    assert output.is_file()
    assert paragraphs["Autonomous Perimeter Monitoring Proposal"] == "Proposal_Title"
    assert paragraphs["Cover"] == "Proposal_CoverTitle"
    assert paragraphs["Solution Module"] == "Proposal_ModuleHeader"
    assert paragraphs["Solution Overview"] == "Proposal_Heading2"
    assert paragraphs["Persistent monitoring."] == "Proposal_BodyText"
    assert paragraphs["EO/IR observation"] == "Proposal_Bullet"
    assert paragraphs["Operational note"] == "Proposal_CalloutText"
    assert paragraphs["Platform overview"] == "Proposal_Caption"
    assert paragraphs["Appendix A"] == "Proposal_Heading1"
    assert paragraphs["REF-1 — Payload specification — Technical library"] == "Proposal_BodyText"
    assert document.sections[0].header.paragraphs[0].text == "Template header"
    assert document.sections[0].footer.paragraphs[0].text == "Template footer"


def test_docx_compiler_preserves_component_and_table_order(tmp_path: Path) -> None:
    template = tmp_path / "master_template.docx"
    output = tmp_path / "proposal.docx"
    _template(template)

    document = DocxDocument(DOCXCompiler(_styles()).compile(_composition(), template, output))
    texts = [paragraph.text for paragraph in document.paragraphs]

    assert texts.index("Cover") < texts.index("Contents") < texts.index("Solution Module") < texts.index("Solution Overview")
    assert len(document.tables) == 2
    assert [cell.text for cell in document.tables[0].rows[0].cells] == ["Requirement", "Response"]
    assert [cell.text for cell in document.tables[0].rows[1].cells] == ["R-1", "Covered"]
    assert [cell.text for cell in document.tables[1].rows[0].cells] == ["R-1", "EO/IR", "Included", "REF-1"]
    assert document.tables[0].rows[0].cells[0].paragraphs[0].style.name == "Proposal_TableHeader"
    assert document.tables[0].rows[1].cells[0].paragraphs[0].style.name == "Proposal_TableCell"
    assert document.tables[1].rows[0].cells[0].paragraphs[0].style.name == "Proposal_RequirementMatrix"


def test_docx_compiler_rejects_missing_master_template_styles(tmp_path: Path) -> None:
    template = tmp_path / "master_template.docx"
    _template(template)
    styles = _styles()
    styles["body"] = BodyStyle(template_style_name="Proposal_Missing")

    with pytest.raises(MissingTemplateStyle, match="Proposal_Missing"):
        DOCXCompiler(styles).compile(_composition(), template, tmp_path / "proposal.docx")

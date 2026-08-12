"""Phase 4B composition-to-template publishing integration coverage."""

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree
import pytest

from proposal_ai_agent.proposal_generation.composition import CompositionEngine
from proposal_ai_agent.proposal_generation.document_plan import ProposalPlan, SectionPlan, SectionRole
from proposal_ai_agent.proposal_generation.docx_compiler import DOCXCompiler
from proposal_ai_agent.proposal_generation.proposal_ir import (
    BulletList,
    Heading,
    KnowledgeReference,
    Paragraph,
    ProposalDocument,
    RequirementMatrix,
    RequirementMatrixEntry,
    Section,
    Table,
    VisualPlaceholder,
)
from proposal_ai_agent.proposal_generation.publishing import (
    MarkdownContentError,
    TemplatePublisher,
    UnsupportedTemplateComponent,
    file_sha256,
    pru_template_semantic_map,
)


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "documents/PRU_T72_Module_Breakdown.docx"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS = {"w": W, "wp": WP, "a": A}


def synthetic_composition(*, unsupported_matrix: bool = False, raw_markdown: bool = False):
    first_blocks = (
        Paragraph(text=(
            "# Distributed coastal monitoring nodes"
            if raw_markdown else
            "Distributed coastal monitoring nodes provide persistent coverage."
        )),
        BulletList(items=("Passive sensing", "Encrypted telemetry")),
        Table(
            headers=("Item", "Specification", "Quantity"),
            rows=(
                ("Sensor", "Multispectral", "4"),
                ("Gateway", "Rugged encrypted", "1"),
                ("Console", "Operator", "2"),
            ),
        ),
        VisualPlaceholder(
            visual_id="VIS-SYNTHETIC",
            description="Coastal sensor network architecture",
            caption="Logical connectivity",
        ),
    )
    if unsupported_matrix:
        first_blocks += (RequirementMatrix(entries=(RequirementMatrixEntry(
            requirement_id="REQ-1", requirement="Synthetic requirement", response="Compliant",
        ),)),)
    architecture = Section(
        section_id="architecture",
        heading=Heading(text="Architecture", level=2),
        blocks=first_blocks,
    )
    package = Section(
        section_id="package",
        heading=Heading(text="Bill of Materials", level=2),
        blocks=(
            Paragraph(text="The deployable package uses modular equipment."),
            Table(
                headers=("S.No", "Item / Equipment", "Specification / Type", "Qty per Kit", "Remarks"),
                rows=(
                    ("1", "Sensor node", "Multispectral", "4", "Field deployable"),
                    ("2", "Gateway", "Encrypted", "1", "Rugged"),
                    ("3", "Console", "Operator station", "2", "Portable"),
                ),
            ),
        ),
    )
    document = ProposalDocument(
        proposal_id="SYN-COAST-001",
        title="Coastal Sensor Network Proposal",
        sections=(
            Section(section_id="cover", heading=Heading(text="Cover", level=1)),
            Section(section_id="module-sensors", heading=Heading(text="Sensor Network", level=1), children=(architecture,)),
            Section(section_id="module-package", heading=Heading(text="Deployment Package", level=1), children=(package,)),
            Section(section_id="references", heading=Heading(text="References", level=1)),
        ),
        references=(KnowledgeReference(
            reference_id="REF-SYN-1", title="Coastal deployment guide", source="Synthetic library",
        ),),
    )
    plan = ProposalPlan(
        proposal_id=document.proposal_id,
        sections=(
            SectionPlan(section_id="cover", role=SectionRole.COVER),
            SectionPlan(section_id="module-sensors", role=SectionRole.MODULE),
            SectionPlan(section_id="architecture", role=SectionRole.BODY),
            SectionPlan(section_id="module-package", role=SectionRole.MODULE),
            SectionPlan(section_id="package", role=SectionRole.BODY),
            SectionPlan(section_id="references", role=SectionRole.REFERENCES),
        ),
    )
    return CompositionEngine().compose(plan, document)


def _part_text(package: ZipFile, name: str) -> str:
    root = etree.fromstring(package.read(name))
    return "".join(root.xpath(".//w:t/text()", namespaces=NS))


def _table_properties(document: Document, index: int) -> tuple[bytes, bytes]:
    table = document.tables[index]._tbl
    return etree.tostring(table.tblPr, method="c14n"), etree.tostring(table.tblGrid, method="c14n")


def _row_fills(document: Document, table_index: int, row_index: int) -> tuple[str, ...]:
    return tuple(
        fill
        for cell in document.tables[table_index].rows[row_index].cells
        for fill in cell._tc.xpath("./w:tcPr/w:shd/@w:fill")
    )


def test_real_template_publisher_binds_composition_and_preserves_package(tmp_path: Path) -> None:
    source_hash = file_sha256(TEMPLATE)
    output = tmp_path / "synthetic.docx"
    result = TemplatePublisher().publish(
        synthetic_composition(), pru_template_semantic_map(TEMPLATE), TEMPLATE, output,
    )
    generated = Document(result)
    source = Document(TEMPLATE)

    assert result == output.resolve()
    assert len(generated.sections) == 17
    assert len(generated.tables) == 3  # cover + variable technical table + BOM
    assert len(generated.tables[1].rows) == 4
    assert len(generated.tables[2].rows) == 4
    assert [cell.text for cell in generated.tables[1].rows[0].cells] == ["Item", "Specification", "Quantity"]
    assert [cell.text for cell in generated.tables[2].rows[-1].cells] == [
        "3", "Console", "Operator station", "2", "Portable",
    ]
    assert _table_properties(generated, 1) == _table_properties(source, 19)
    assert _table_properties(generated, 2) == _table_properties(source, 13)
    assert tuple(_row_fills(generated, 2, row) for row in range(1, 4)) == tuple(
        _row_fills(source, 13, row) for row in range(1, 4)
    )

    paragraphs = {paragraph.text: paragraph for paragraph in generated.paragraphs}
    assert paragraphs["Architecture\t"].style.name == "Heading 1"
    assert paragraphs["Distributed coastal monitoring nodes provide persistent coverage."].style.name == "Body Text"
    for item in ("Passive sensing", "Encrypted telemetry"):
        paragraph = paragraphs[item]
        assert paragraph._p.xpath("./w:pPr/w:numPr/w:numId/@w:val") == ["1"]
        assert paragraph._p.xpath("./w:pPr/w:numPr/w:ilvl/@w:val") == ["3"]

    with ZipFile(result) as package:
        names = package.namelist()
        document_xml = etree.fromstring(package.read("word/document.xml"))
        headers = [name for name in names if name.startswith("word/header") and name.endswith(".xml")]
        footers = [name for name in names if name.startswith("word/footer") and name.endswith(".xml")]
        media = [name for name in names if name.startswith("word/media/")]
        assert len(headers) == len(footers) == 5
        assert len(media) == 8
        assert all(_part_text(package, name).count("Coastal Sensor Network Proposal") == 2 for name in headers)
        assert all(
            {value.strip() for value in etree.fromstring(package.read(name)).xpath(".//w:instrText/text()", namespaces=NS)}
            >= {"PAGE", "NUMPAGES"}
            for name in footers
        )
        banner_drawings = document_xml.xpath(
            './/w:drawing[.//wp:extent[@cy="393700"]][.//a:solidFill/a:srgbClr[@val="E96F30"]]',
            namespaces=NS,
        )
        assert len(banner_drawings) == 2
        assert all(drawing.xpath('.//a:ln/a:solidFill/a:srgbClr/@val', namespaces=NS) == ["C2541B"] for drawing in banner_drawings)
        body_nodes = document_xml.xpath("./w:body/*", namespaces=NS)
        ordered = ["".join(node.xpath(".//w:t/text()", namespaces=NS)) for node in body_nodes]
        positions = [
            next(index for index, text in enumerate(ordered) if marker in text)
            for marker in (
                "Sensor Network", "Architecture", "Distributed coastal monitoring",
                "ItemSpecificationQuantity", "Deployment Package", "Bill of Materials",
                "S.NoItem / EquipmentSpecification / TypeQty per KitRemarks", "References", "REF-SYN-1",
            )
        ]
        assert positions == sorted(positions)
        complete_text = "\n".join(_part_text(package, name) for name in ["word/document.xml", *headers])
        for stale in (
            "T-72", "Drive-by-Wire", "SENTINEL ACUS", "17 July 2026",
            "Draft — For Internal Review", "AI INTEGRATION",
        ):
            assert stale not in complete_text
        assert not any(line.lstrip().startswith(("# ", "- ", "* ")) for line in complete_text.splitlines())

    assert file_sha256(TEMPLATE) == source_hash


def test_requirement_matrix_fails_with_selected_template_and_section_diagnostic(tmp_path: Path) -> None:
    with pytest.raises(
        UnsupportedTemplateComponent,
        match=r"component_type=requirement_matrix section_id=architecture.*no verified compliance",
    ):
        TemplatePublisher().publish(
            synthetic_composition(unsupported_matrix=True),
            pru_template_semantic_map(TEMPLATE), TEMPLATE, tmp_path / "unsupported.docx",
        )
    assert not (tmp_path / "unsupported.docx").exists()


def test_raw_markdown_fails_before_template_copy_is_published(tmp_path: Path) -> None:
    output = tmp_path / "markdown.docx"
    with pytest.raises(MarkdownContentError, match="raw Markdown"):
        TemplatePublisher().publish(
            synthetic_composition(raw_markdown=True),
            pru_template_semantic_map(TEMPLATE), TEMPLATE, output,
        )
    assert not output.exists()


def test_docx_compiler_delegates_pru_template_to_template_publisher(tmp_path: Path, monkeypatch) -> None:
    output = tmp_path / "delegated.docx"
    calls = []

    def publish(self, composition, template_map, source_template, output_path):
        calls.append((composition, template_map, Path(source_template), Path(output_path)))
        output.write_bytes(b"delegated")
        return output

    monkeypatch.setattr(TemplatePublisher, "publish", publish)
    result = DOCXCompiler({}).compile(synthetic_composition(), TEMPLATE, output)

    assert result == output
    assert len(calls) == 1
    assert calls[0][2] == TEMPLATE
    assert calls[0][1].template_sha256 == file_sha256(TEMPLATE)

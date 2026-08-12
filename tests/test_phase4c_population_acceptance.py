"""Acceptance checks for deterministic certified population of the real PRU."""

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import pytest

from proposal_ai_agent.proposal_generation.publishing import (
    BindingAuthorizationError,
    CertifiedContracts,
    CertifiedTemplateBinding,
    file_sha256,
)
from scripts.run_phase4c_population_acceptance import (
    CONTRACTS,
    EXPECTED_SHA,
    SOURCE,
    _validate_output,
    run_acceptance,
)


ROOT = Path(__file__).resolve().parents[1]
ARTIFACT = ROOT / "output/phase4c_population_acceptance.docx"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _property_bytes(node) -> tuple[bytes, ...]:
    properties = node.xpath(".//w:pPr | .//w:rPr | .//w:tcPr | .//w:trPr")
    return tuple(etree.tostring(item, method="c14n", exclusive=True) for item in properties)


def test_acceptance_generation_uses_copy_and_preserves_source(tmp_path: Path) -> None:
    before = file_sha256(SOURCE)
    output = run_acceptance(tmp_path / "acceptance.docx")
    assert output != SOURCE
    assert output.is_file()
    assert file_sha256(SOURCE) == before == EXPECTED_SHA
    _validate_output(SOURCE, output)


def test_acceptance_artifact_contains_deterministic_synthetic_values() -> None:
    document = Document(ARTIFACT)
    body_text = "".join(document.element.body.xpath(".//w:t/text()"))
    header_text = "".join(
        "".join(part.element.xpath(".//w:t/text()"))
        for part in document.part.package.parts
        if str(part.partname).startswith("/word/header")
    )
    assert "BDIL SKYSHIELD-X Acceptance Demonstrator" in body_text
    assert "Autonomous Multi-Sensor Defence Platform" in body_text
    assert "Synthetic Quantum Radar Controller" in body_text
    assert "Certified Cloned Sensor Node" in body_text
    assert "Duplicate-Table Isolation Beacon" in body_text
    assert header_text.count("BDIL SKYSHIELD-X — CONTROLLED ACCEPTANCE TEST") == 10
    assert "#" not in body_text + header_text
    assert "**" not in body_text + header_text


def test_run_paragraph_cell_and_row_formatting_survives_population() -> None:
    source = Document(SOURCE)
    output = Document(ARTIFACT)
    source_nodes = tuple(source.element.body.iterchildren())
    output_nodes = tuple(output.element.body.iterchildren())
    for body_index in (4, 13, 127):
        assert _property_bytes(output_nodes[body_index]) == _property_bytes(source_nodes[body_index])

    for table_index, rows in ((0, (1, 5)), (1, (0,)), (12, (1, 2)), (13, (1, 2)), (15, (1,))):
        for row_index in rows:
            assert _property_bytes(output.tables[table_index].rows[row_index]._tr) == _property_bytes(
                source.tables[table_index].rows[row_index]._tr
            )
    assert len(output.tables[13].rows) == len(source.tables[13].rows) + 1
    assert _property_bytes(output.tables[13].rows[6]._tr) == _property_bytes(source.tables[13].rows[2]._tr)


def test_every_certified_static_table_cell_remains_unchanged() -> None:
    contracts = CertifiedContracts.load(CONTRACTS)
    source = Document(SOURCE)
    output = Document(ARTIFACT)
    for binding in contracts.template_map["static_table_cells"]:
        locator = binding["locator"]
        source_row = source.tables[locator["table_index"]]._tbl.findall(qn("w:tr"))[locator["row_index"]]
        output_row = output.tables[locator["table_index"]]._tbl.findall(qn("w:tr"))[locator["row_index"]]
        source_cell = source_row.findall(qn("w:tc"))[locator["physical_cell_index"]]
        output_cell = output_row.findall(qn("w:tc"))[locator["physical_cell_index"]]
        assert etree.tostring(output_cell, method="c14n") == etree.tostring(source_cell, method="c14n")


def test_acceptance_binding_rejects_markup_and_uncertified_mutation(tmp_path: Path) -> None:
    contracts = CertifiedContracts.load(CONTRACTS)
    binding = CertifiedTemplateBinding(SOURCE, contracts, tmp_path / "copy.docx")
    with pytest.raises(BindingAuthorizationError):
        binding.replace_dynamic_text("m1_integration.text_slot_01", "## forbidden")
    with pytest.raises(BindingAuthorizationError):
        binding.replace_dynamic_text("module_1", "uncertified")
    with pytest.raises(BindingAuthorizationError):
        binding.replace_dynamic_cell(
            "bdil_pru_v1.table.00.cover.document_control.r0.c0", "static mutation"
        )


def test_media_and_computed_fields_are_byte_and_instruction_identical() -> None:
    with ZipFile(SOURCE) as source, ZipFile(ARTIFACT) as output:
        media = [name for name in source.namelist() if name.startswith("word/media/")]
        assert len(media) == 8
        assert all(source.read(name) == output.read(name) for name in media)
        for index in range(1, 6):
            source_footer = etree.fromstring(source.read(f"word/footer{index}.xml"))
            output_footer = etree.fromstring(output.read(f"word/footer{index}.xml"))
            source_fields = source_footer.xpath(".//w:instrText/text()", namespaces={"w": W})
            output_fields = output_footer.xpath(".//w:instrText/text()", namespaces={"w": W})
            assert output_fields == source_fields

"""Build the crisp Phase 4C.1.5 artifact using certified mutations only."""

from __future__ import annotations

from hashlib import sha256
import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

from proposal_ai_agent.proposal_generation.publishing import (
    CertifiedContracts,
    CertifiedTemplateBinding,
    file_sha256,
)
ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "documents/PRU_T72_Module_Breakdown.docx"
CONTRACTS = ROOT / "templates/bdil_pru_v1"
EXPECTED_SHA = "e81d6c34362fb0e7d17b8a82e5e756ad696f98d61e90100e3ca5b7bec7182d55"
OUTPUT = ROOT / "output/phase4c_population_acceptance_crisp.docx"
VISUAL_OUTPUT = ROOT / "output/phase4c_visual_acceptance_crisp"
PACKAGE_DIFF = VISUAL_OUTPUT / "package_diff.json"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def _cell_id(table: int, owner: str, purpose: str, row: int, column: int) -> str:
    return f"bdil_pru_v1.table.{table:02d}.{owner}.{purpose}.r{row}.c{column}"


def _populate_existing_row(
    binding: CertifiedTemplateBinding,
    table: int,
    owner: str,
    purpose: str,
    row: int,
    values: tuple[str, ...],
) -> None:
    for column, value in enumerate(values):
        binding.replace_dynamic_cell(
            _cell_id(table, owner, purpose, row, column), value
        )


def _part_digest(package: ZipFile, name: str) -> str:
    return sha256(package.read(name)).hexdigest()


def _display_path(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(ROOT))
    except ValueError:
        return str(path.resolve())


def package_diff(source_path: Path, output_path: Path) -> dict[str, object]:
    with ZipFile(source_path) as source, ZipFile(output_path) as output:
        source_names = set(source.namelist())
        output_names = set(output.namelist())
        changed = sorted(
            name for name in source_names & output_names
            if _part_digest(source, name) != _part_digest(output, name)
        )
        return {
            "source": _display_path(source_path),
            "output": _display_path(output_path),
            "changed_parts": changed,
            "removed_parts": sorted(source_names - output_names),
            "added_parts": sorted(output_names - source_names),
            "expected_changed_parts": [
                "word/document.xml", "word/_rels/document.xml.rels",
                *[f"word/header{index}.xml" for index in range(1, 6)],
            ],
            "expected_removed_parts": [f"word/media/image{index}.jpeg" for index in range(2, 9)],
            "unexplained_changed_parts": sorted(set(changed) - {
                "word/document.xml", "word/_rels/document.xml.rels",
                *{f"word/header{index}.xml" for index in range(1, 6)},
            }),
            "unexplained_removed_parts": sorted(
                (source_names - output_names) - {f"word/media/image{index}.jpeg" for index in range(2, 9)}
            ),
        }


def _validate_hardened_output(output: Path, contracts: CertifiedContracts) -> None:
    document = Document(output)
    source = Document(SOURCE)
    assert len(document.sections) == len(source.sections) == 17
    assert len(document.tables) == len(source.tables) == 20
    assert len(document.tables[13].rows) == len(source.tables[13].rows) + 1
    body_nodes = tuple(document.element.body.iterchildren())
    locations = contracts.schema["manual_diagram_locations"]
    assert len(locations) == 7
    for location in locations:
        node = body_nodes[location["locator"]["body_node_index"]]
        assert "".join(node.xpath(".//w:t/text()")) == "[ INSERT DIAGRAM ]"
        assert not node.xpath(".//w:drawing")
    all_text = "".join(document.element.body.xpath(".//w:t/text()"))
    assert "#" not in all_text and "**" not in all_text
    with ZipFile(SOURCE) as source_package, ZipFile(output) as output_package:
        assert output_package.read("word/media/image1.png") == source_package.read("word/media/image1.png")
        assert not any(
            name.startswith("word/media/") and name.endswith(".jpeg")
            for name in output_package.namelist()
        )
        source_xml = etree.fromstring(source_package.read("word/document.xml"))
        output_xml = etree.fromstring(output_package.read("word/document.xml"))
        source_sections = source_xml.xpath(".//w:sectPr", namespaces=NS)
        output_sections = output_xml.xpath(".//w:sectPr", namespaces=NS)
        assert [etree.tostring(x, method="c14n") for x in output_sections] == [
            etree.tostring(x, method="c14n") for x in source_sections
        ]
        for name in (
            "word/styles.xml", "word/theme/theme1.xml", "word/numbering.xml",
            "word/settings.xml", "word/footer1.xml", "word/footer2.xml",
            "word/footer3.xml", "word/footer4.xml", "word/footer5.xml",
        ):
            assert output_package.read(name) == source_package.read(name)


def run_hardening(output: Path = OUTPUT) -> Path:
    source_before = file_sha256(SOURCE)
    if source_before != EXPECTED_SHA:
        raise RuntimeError("source template does not match certification")
    contracts = CertifiedContracts.load(CONTRACTS)
    binding = CertifiedTemplateBinding(SOURCE, contracts, working_copy=output)

    # Concise values reduce avoidable wrapping while template formatting remains untouched.
    binding.replace_dynamic_text("cover_document_title", "SKYSHIELD-X ACCEPTANCE PRU")
    binding.replace_dynamic_text(
        "m1_integration.text_slot_01",
        "Certified body-text population preserves template formatting.",
    )
    binding.replace_dynamic_text(
        "m5_integration.text_slot_18",
        "SKYSHIELD-X uses synthetic sensor inputs for controlled acceptance.",
    )
    for index in range(1, 6):
        binding.replace_dynamic_text(
            f"header_project_text_{index}", "SKYSHIELD-X — ACCEPTANCE PRU"
        )
    binding.replace_dynamic_cell(
        _cell_id(0, "cover", "document_control", 1, 1),
        "Multi-Sensor Defence Demonstrator",
    )
    binding.replace_dynamic_cell(
        _cell_id(0, "cover", "document_control", 5, 1), "ACCEPTANCE TEST"
    )
    _populate_existing_row(
        binding, 1, "m1_power", "power_consumption", 0,
        ("1", "Synthetic Sensor Unit", "111 W"),
    )
    _populate_existing_row(
        binding, 12, "m5_power", "power_placeholder", 1,
        ("1", "Fusion Processor", "2", "48", "222 W", "Test fixture"),
    )
    _populate_existing_row(
        binding, 12, "m5_power", "power_placeholder", 2,
        ("2", "Optical Array", "4", "48", "144 W", "Synthetic"),
    )
    clone = binding.insert_certified_row(
        "bdil_pru_v1.table.13.m5_bom.bom_placeholder.row_prototype.2",
        insert_after_row_index=5,
    )
    binding.populate_certified_cloned_row(
        clone, ("6", "Cloned Sensor Node", "Type CX-6", "1", "Certified clone")
    )
    _populate_existing_row(
        binding, 13, "m5_bom", "bom_placeholder", 1,
        ("1", "Synthetic Radar Controller", "Type QR-1", "1", "Test fixture"),
    )
    _populate_existing_row(
        binding, 13, "m5_bom", "bom_placeholder", 2,
        ("2", "Telemetry Emulator", "Type TE-2", "2", "Synthetic"),
    )
    _populate_existing_row(
        binding, 15, "m6_bom", "bom_placeholder", 1,
        ("1", "Isolation Beacon", "Type IB-1", "1", "M6 only"),
    )
    for location in contracts.schema["manual_diagram_locations"]:
        binding.replace_certified_manual_diagram(location["location_id"])

    result = binding.save(output)
    if file_sha256(SOURCE) != source_before:
        raise RuntimeError("source template changed during visual hardening")
    _validate_hardened_output(result, contracts)
    diff = package_diff(SOURCE, result)
    if diff["unexplained_changed_parts"] or diff["unexplained_removed_parts"]:
        raise RuntimeError("hardened package contains unexplained OPC changes")
    if result.resolve() == OUTPUT.resolve():
        VISUAL_OUTPUT.mkdir(parents=True, exist_ok=True)
        PACKAGE_DIFF.write_text(json.dumps(diff, indent=2) + "\n", encoding="utf-8")
    return result


if __name__ == "__main__":
    generated = run_hardening()
    print(f"Created: {generated}")
    print(f"Source SHA-256 before/after: {EXPECTED_SHA}")
    print(f"Package diff: {PACKAGE_DIFF}")

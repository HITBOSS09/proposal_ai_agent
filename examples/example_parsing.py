"""
Example: Complete Document Parsing Workflow

This module demonstrates the end-to-end parsing process with a realistic
proposal document. It shows:
1. Document structure before and after parsing
2. How the parser handles hierarchy and nesting
3. How metadata is propagated
4. How order_index preserves reading order
"""

from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Pt

from proposal_ai_agent.ingestion.parser import DocumentParser
from proposal_ai_agent.ingestion.models import Table


def create_example_proposal() -> Path:
    """
    Create a realistic proposal document for parsing demonstration.
    
    Document structure:
    - Preamble (root level)
    - Executive Summary (H1)
      - Overview (H2)
      - Key Metrics (H2 with table)
    - Implementation Plan (H1)
      - Timeline (H2)
      - Resource Allocation (H2)
    - Appendix (H1)
    """
    proposal_path = Path("/tmp/example_proposal.docx")
    
    doc = DocxDocument()
    
    # === ROOT LEVEL PREAMBLE ===
    doc.add_paragraph(
        "This proposal outlines a comprehensive digital transformation initiative "
        "designed to modernize our operations and improve customer engagement.",
        style="Normal"
    )
    
    # === SECTION 1: EXECUTIVE SUMMARY ===
    doc.add_paragraph("Executive Summary", style="Heading 1")
    
    # Subsection 1.1: Overview
    doc.add_paragraph("Overview", style="Heading 2")
    doc.add_paragraph(
        "The initiative encompasses three core pillars: infrastructure modernization, "
        "customer experience improvement, and operational efficiency gains.",
        style="Normal"
    )
    doc.add_paragraph(
        "Expected outcomes include 40% reduction in processing time and 25% improvement "
        "in customer satisfaction scores.",
        style="Normal"
    )
    
    # Subsection 1.2: Key Metrics
    doc.add_paragraph("Key Performance Indicators", style="Heading 2")
    doc.add_paragraph(
        "The following table outlines our projected KPIs:",
        style="Normal"
    )
    
    # Add metrics table
    table = doc.add_table(rows=5, cols=3)
    table.style = "Light Grid Accent 1"
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Metric"
    header_cells[1].text = "Current State"
    header_cells[2].text = "Target (Year 1)"
    
    # Data rows
    metrics = [
        ("Processing Time (hours)", "8", "4.8"),
        ("Customer Satisfaction (%)", "72", "90"),
        ("System Uptime (%)", "99.2", "99.95"),
    ]
    
    for idx, (metric, current, target) in enumerate(metrics, start=1):
        row_cells = table.rows[idx].cells
        row_cells[0].text = metric
        row_cells[1].text = current
        row_cells[2].text = target
    
    # === SECTION 2: IMPLEMENTATION PLAN ===
    doc.add_paragraph("Implementation Plan", style="Heading 1")
    
    # Subsection 2.1: Timeline
    doc.add_paragraph("Project Timeline", style="Heading 2")
    doc.add_paragraph(
        "The implementation will be executed in four phases over a 12-month period:",
        style="Normal"
    )
    
    # Add timeline content
    doc.add_paragraph("Phase 1: Planning & Setup (Months 1-2)", style="Normal")
    doc.add_paragraph("Infrastructure assessment and team allocation", style="Normal")
    doc.add_paragraph(
        "Phase 2: Core Infrastructure (Months 3-6)",
        style="Normal"
    )
    doc.add_paragraph(
        "Database migration and API modernization",
        style="Normal"
    )
    doc.add_paragraph("Phase 3: Integration & Testing (Months 7-10)", style="Normal")
    doc.add_paragraph("System integration and comprehensive testing", style="Normal")
    doc.add_paragraph("Phase 4: Rollout & Optimization (Months 11-12)", style="Normal")
    doc.add_paragraph("Production deployment and performance tuning", style="Normal")
    
    # Subsection 2.2: Resource Allocation
    doc.add_paragraph("Resource Allocation", style="Heading 2")
    
    # Budget table
    budget_table = doc.add_table(rows=5, cols=2)
    budget_table.style = "Light Grid Accent 1"
    
    budget_cells = budget_table.rows[0].cells
    budget_cells[0].text = "Resource Category"
    budget_cells[1].text = "Budget ($M)"
    
    budget_data = [
        ("Engineering & Development", "2.5"),
        ("Infrastructure & Operations", "1.2"),
        ("Project Management & Consulting", "0.8"),
    ]
    
    for idx, (category, budget) in enumerate(budget_data, start=1):
        row_cells = budget_table.rows[idx].cells
        row_cells[0].text = category
        row_cells[1].text = budget
    
    # === SECTION 3: APPENDIX ===
    doc.add_paragraph("Appendix", style="Heading 1")
    doc.add_paragraph("Additional documentation and reference materials are available upon request.", style="Normal")
    
    # Save the document
    doc.save(proposal_path)
    print(f"✓ Created example proposal: {proposal_path}")
    
    return proposal_path


def print_document_structure(doc, indent=0):
    """
    Pretty-print the internal Document structure.
    
    Args:
        doc: Internal Document object
        indent: Current indentation level
    """
    spaces = "  " * indent
    
    print(f"\n{spaces}📄 DOCUMENT")
    print(f"{spaces}├─ Title: {doc.title}")
    print(f"{spaces}├─ Document ID: {doc.document_id}")
    print(f"{spaces}├─ Author: {doc.metadata.author or 'N/A'}")
    print(f"{spaces}├─ Source: {doc.metadata.source}")
    
    if doc.elements:
        print(f"{spaces}├─ Root Elements: {len(doc.elements)}")
        for elem in doc.elements:
            print(f"{spaces}│  ├─ [{elem.metadata.order_index}] {type(elem).__name__.lower()}")
            if hasattr(elem, 'content'):
                preview = elem.content[:60].replace('\n', ' ')
                print(f"{spaces}│  │  └─ \"{preview}...\"")
    
    if doc.sections:
        print(f"{spaces}└─ Sections: {len(doc.sections)}")
        for section_idx, section in enumerate(doc.sections):
            is_last = section_idx == len(doc.sections) - 1
            section_prefix = "    " if is_last else "├─"
            print_section_structure(section, indent + 1, 0)


def print_section_structure(section, indent, parent_order):
    """
    Recursively print section structure with elements.
    
    Args:
        section: Section object to print
        indent: Current indentation level
        parent_order: Order index of parent (for context)
    """
    spaces = "  " * indent
    print(f"{spaces}📂 [{section.section_level}] {section.heading}")
    
    if section.elements:
        print(f"{spaces}├─ Elements: {len(section.elements)}")
        for elem_idx, elem in enumerate(section.elements):
            is_last_elem = elem_idx == len(section.elements) - 1 and not section.subsections
            elem_prefix = "└─" if is_last_elem else "├─"
            print(f"{spaces}{elem_prefix} [{elem.metadata.order_index}] {type(elem).__name__.lower()}")
            
            if hasattr(elem, 'content'):
                preview = elem.content[:55].replace('\n', ' ')
                print(f"{spaces}  └─ \"{preview}...\"")
            elif isinstance(elem, Table):
                print(f"{spaces}  └─ Table ({len(elem.rows)} rows, {len(elem.rows[0].cells)} cols)")
    
    if section.subsections:
        print(f"{spaces}└─ Subsections: {len(section.subsections)}")
        for sub_idx, subsection in enumerate(section.subsections):
            print_section_structure(subsection, indent + 1, parent_order + len(section.elements))


def main():
    """
    Main execution: Create and parse example proposal.
    """
    print("=" * 80)
    print("DOCUMENT PARSER - EXAMPLE EXECUTION")
    print("=" * 80)
    
    # === STEP 1: CREATE EXAMPLE DOCUMENT ===
    print("\n[STEP 1] Creating example proposal document...")
    docx_path = create_example_proposal()
    
    # === STEP 2: LOAD DOCUMENT ===
    print("\n[STEP 2] Loading DOCX document...")
    docx_doc = DocxDocument(docx_path)
    print(f"✓ Loaded {len(docx_doc.paragraphs)} paragraphs, {len(docx_doc.tables)} tables")
    
    # === STEP 3: PARSE DOCUMENT ===
    print("\n[STEP 3] Parsing document to internal model...")
    parser = DocumentParser()
    parsed_doc = parser.parse(
        docx_doc,
        docx_path,
        source_document="Proposal 2024",
        author="Strategic Planning Team"
    )
    print(f"✓ Parsing complete")
    
    # === STEP 4: DISPLAY STRUCTURE ===
    print("\n[STEP 4] Document Structure:")
    print_document_structure(parsed_doc)
    
    # === STEP 5: DETAILED ANALYSIS ===
    print("\n\n[STEP 5] Detailed Analysis:")
    print(f"  Total Sections (top-level): {len(parsed_doc.sections)}")
    
    total_elements = len(parsed_doc.elements)
    def count_elements(sections):
        nonlocal total_elements
        for section in sections:
            total_elements += len(section.elements)
            count_elements(section.subsections)
    
    count_elements(parsed_doc.sections)
    print(f"  Total Elements (all): {total_elements}")
    
    # Count by type
    element_types = {}
    def count_by_type(sections):
        for section in sections:
            for elem in section.elements:
                elem_type = type(elem).__name__.lower()
                element_types[elem_type] = element_types.get(elem_type, 0) + 1
            count_by_type(section.subsections)
    
    count_by_type(parsed_doc.sections)
    for elem_type in parsed_doc.elements:
        elem_type_str = type(elem_type).__name__.lower()
        element_types[elem_type_str] = element_types.get(elem_type_str, 0) + 1
    
    print(f"  Element Breakdown:")
    for elem_type, count in sorted(element_types.items()):
        print(f"    - {elem_type.title()}: {count}")
    
    # === STEP 6: SAMPLE ELEMENT INSPECTION ===
    print("\n\n[STEP 6] Sample Element Inspection:")
    
    if parsed_doc.sections:
        section1 = parsed_doc.sections[0]
        print(f"\n  First Section: {section1.heading}")
        print(f"  └─ Section Level: {section1.section_level}")
        print(f"  └─ Direct Elements: {len(section1.elements)}")
        print(f"  └─ Subsections: {len(section1.subsections)}")
        
        if section1.elements:
            first_elem = section1.elements[0]
            print(f"\n  First Element in Section:")
            print(f"  ├─ Type: {type(first_elem).__name__.lower()}")
            print(f"  ├─ Order Index: {first_elem.metadata.order_index}")
            print(f"  ├─ Document ID: {first_elem.metadata.document_id}")
            print(f"  ├─ Section ID: {first_elem.metadata.section_id}")
            print(f"  ├─ Source File: {first_elem.metadata.source_file}")
            
            if hasattr(first_elem, 'content'):
                print(f"  └─ Content (first 80 chars): \"{first_elem.content[:80]}...\"")
        
        if section1.subsections:
            sub = section1.subsections[0]
            print(f"\n  First Subsection: {sub.heading}")
            print(f"  └─ Section Level: {sub.section_level}")
            print(f"  └─ Elements: {len(sub.elements)}")
            
            if sub.elements and isinstance(sub.elements[0], Table):
                table = sub.elements[0]
                print(f"\n  Sample Table:")
                print(f"  ├─ Rows: {len(table.rows)}")
                print(f"  ├─ Columns: {len(table.rows[0].cells) if table.rows else 0}")
                print(f"  └─ Header Row:")
                if table.rows:
                    for cell in table.rows[0].cells:
                        print(f"     - \"{cell.content[:20]}...\"")
    
    print("\n" + "=" * 80)
    print("✓ PARSING COMPLETE")
    print("=" * 80)
    
    return parsed_doc


if __name__ == "__main__":
    doc = main()

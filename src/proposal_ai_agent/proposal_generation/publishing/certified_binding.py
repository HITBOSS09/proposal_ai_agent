"""Fail-closed bindings for the version-pinned BDIL PRU template.

This module authorizes only certified content replacement and prototype-row
cloning. It deliberately does not publish proposals or interpret model output.
"""

from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from enum import Enum
from hashlib import sha256
import json
import os
from pathlib import Path
import tempfile
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from .openxml_editor import (
    LocatorMismatch,
    OpenXmlEditor,
    TemplateIntegrityError,
    UnsafeOpenXmlOperation,
    copy_template,
    file_sha256,
)


MC_ALTERNATE_CONTENT = (
    "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent"
)


class StructuralOperation(str, Enum):
    KEEP = "KEEP"
    REPLACE_CONTENT = "REPLACE_CONTENT"
    INSERT_BEFORE = "INSERT_BEFORE"
    INSERT_AFTER = "INSERT_AFTER"
    REMOVE = "REMOVE"
    REORDER = "REORDER"


class BindingAuthorizationError(UnsafeOpenXmlOperation):
    """The requested mutation is not authorized by the certified policy."""


@dataclass(frozen=True)
class CertifiedClonedRow:
    """An in-memory row derived from one exact certified prototype."""

    prototype_id: str
    table_id: str
    element: Any


def _digest(node: Any) -> str | None:
    if node is None:
        return None
    return sha256(etree.tostring(node, method="c14n")).hexdigest()


@dataclass(frozen=True)
class CertifiedContracts:
    root: Path
    schema: dict[str, Any]
    template_map: dict[str, Any]
    policy: dict[str, Any]
    manifest: dict[str, Any]

    @classmethod
    def load(cls, root: str | Path) -> "CertifiedContracts":
        location = Path(root)
        payloads = {
            name: json.loads((location / filename).read_text(encoding="utf-8"))
            for name, filename in {
                "schema": "template_schema.json",
                "template_map": "template_map.json",
                "policy": "template_policy.json",
                "manifest": "manifest.json",
            }.items()
        }
        template_ids = {
            payloads[name]["template_id"] for name in payloads
        }
        if template_ids != {"BDIL_PRU_V1"}:
            raise TemplateIntegrityError("certification artifacts disagree on template identity")
        expected = payloads["schema"]["source_template"]["sha256"]
        if payloads["template_map"]["template_sha256"] != expected:
            raise TemplateIntegrityError("schema and map template hashes disagree")
        if payloads["manifest"]["source_sha256"] != expected:
            raise TemplateIntegrityError("manifest and schema template hashes disagree")
        schema_diagrams = {
            item["location_id"] for item in payloads["schema"].get("manual_diagram_locations", [])
        }
        map_diagrams = {
            item["location_id"] for item in payloads["template_map"].get("manual_diagram_locations", [])
        }
        policy_diagrams = set(
            payloads["policy"]["image_policy"].get("PROJECT_SPECIFIC_TECHNICAL_VISUAL", [])
        )
        if schema_diagrams != map_diagrams or schema_diagrams != policy_diagrams:
            raise TemplateIntegrityError("manual-diagram certifications disagree")
        return cls(root=location, **payloads)

    @property
    def template_sha256(self) -> str:
        return self.schema["source_template"]["sha256"]

    def validate_structural_operation(
        self,
        operation: StructuralOperation,
        *,
        target_id: str | None = None,
        boundary_id: str | None = None,
        explicit_user_override: bool = False,
    ) -> None:
        if operation is StructuralOperation.KEEP:
            return
        if operation is StructuralOperation.REPLACE_CONTENT:
            editable = {
                item["field_id"] for item in self.schema["dynamic_text_fields"]
            } | {
                item["binding_id"]
                for item in self.template_map["dynamic_table_cell_bindings"]
            }
            if target_id not in editable:
                raise BindingAuthorizationError("target is not a certified dynamic node")
            return
        if operation in {StructuralOperation.INSERT_BEFORE, StructuralOperation.INSERT_AFTER}:
            allowed = set(
                self.policy["operations"][operation.value]["allowed_boundaries"]
            )
            if not explicit_user_override:
                raise BindingAuthorizationError("insert requires an explicit user override")
            if boundary_id not in allowed:
                raise BindingAuthorizationError("insert boundary is not USER_INSERTABLE")
            return
        if operation in {StructuralOperation.REMOVE, StructuralOperation.REORDER}:
            raise BindingAuthorizationError(f"{operation.value} has no certified target")
        raise BindingAuthorizationError("unsupported structural operation")


class CertifiedTemplateBinding:
    """Resolve and mutate only exact nodes named by the certified contracts."""

    def __init__(
        self,
        source_template: str | Path,
        contracts: CertifiedContracts,
        working_copy: str | Path | None = None,
    ) -> None:
        self.source_template = Path(source_template).resolve()
        if file_sha256(self.source_template) != contracts.template_sha256:
            raise TemplateIntegrityError("source template SHA-256 does not match certification")
        self.source_hash = contracts.template_sha256
        self.contracts = contracts
        self.working_copy = (
            copy_template(self.source_template, working_copy)
            if working_copy is not None
            else self.source_template
        )
        self.document = Document(self.working_copy)
        self._tables = {item["table_id"]: item for item in contracts.template_map["tables"]}
        self._dynamic_cells = {
            item["binding_id"]: item
            for item in contracts.template_map["dynamic_table_cell_bindings"]
        }
        self._static_cells = {
            item["binding_id"] for item in contracts.template_map["static_table_cells"]
        }
        self._prototypes = {
            item["prototype_id"]: item
            for item in contracts.template_map["repeatable_row_prototypes"]
        }
        self._dynamic_text = {
            item["field_id"]: item for item in contracts.schema["dynamic_text_fields"]
        }
        self._manual_diagrams = {
            item["location_id"]: item
            for item in contracts.schema.get("manual_diagram_locations", [])
        }
        if len(self._tables) != 20:
            raise TemplateIntegrityError("all 20 certified table identities are required")
        self._validate_all_table_locators()

    @staticmethod
    def _validate_plain_text(value: str) -> None:
        if not isinstance(value, str) or not value.strip():
            raise BindingAuthorizationError("certified replacement must be non-empty text")
        if "#" in value or "**" in value:
            raise BindingAuthorizationError("Markdown presentation markers are forbidden")

    def _story_root(self, story: str):
        if story == "word/document.xml":
            return self.document.element
        matches = [
            part for part in self.document.part.package.parts
            if str(part.partname).lstrip("/") == story
        ]
        if len(matches) != 1 or not hasattr(matches[0], "element"):
            raise LocatorMismatch(f"story part did not resolve uniquely: {story}")
        return matches[0].element

    @staticmethod
    def _drawing_by_name(root, name: str):
        matches = root.xpath(f'.//w:drawing[.//wp:docPr[@name="{name}"]]')
        if len(matches) != 1:
            raise LocatorMismatch(f"drawing did not resolve uniquely: {name}")
        return matches[0]

    def replace_dynamic_text(self, field_id: str, replacement: str) -> int:
        """Replace one exact certified text field without replacing its containers."""

        self._validate_plain_text(replacement)
        self.contracts.validate_structural_operation(
            StructuralOperation.REPLACE_CONTENT, target_id=field_id
        )
        try:
            field = self._dynamic_text[field_id]
        except KeyError as error:
            raise BindingAuthorizationError("text field is not certified") from error
        locator = field["locator"]
        expected = locator["verified_text"]
        if locator["element_kind"] == "paragraph":
            nodes = tuple(self.document.element.body.iterchildren())
            try:
                paragraph = nodes[locator["body_node_index"]]
            except IndexError as error:
                raise LocatorMismatch("paragraph locator is out of range") from error
            if _digest(paragraph) != locator["structural_xml_signature"]:
                raise LocatorMismatch("paragraph structural signature mismatch")
            if expected not in "".join(paragraph.xpath(".//w:t/text()")):
                raise LocatorMismatch("paragraph verified text mismatch")
            return OpenXmlEditor._replace_in_container(paragraph, expected, replacement)
        if locator["element_kind"] == "DrawingML_and_VML_textbox":
            root = self._story_root(locator["story"])
            if locator.get("body_node_index") is not None:
                nodes = tuple(self.document.element.body.iterchildren())
                try:
                    body_node = nodes[locator["body_node_index"]]
                except IndexError as error:
                    raise LocatorMismatch("drawing body locator is out of range") from error
                if _digest(body_node) != locator["structural_xml_signature"]:
                    raise LocatorMismatch("drawing structural signature mismatch")
                root = body_node
            drawing = self._drawing_by_name(root, locator["drawing_name"])
            scope = drawing
            for ancestor in drawing.iterancestors():
                if ancestor.tag == MC_ALTERNATE_CONTENT:
                    scope = ancestor
                    break
            text_boxes = list(scope.iter(qn("w:txbxContent")))
            if not text_boxes or any(
                expected not in "".join(node.text or "" for node in box.iter(qn("w:t")))
                for box in text_boxes
            ):
                raise LocatorMismatch("DrawingML/VML verified text mismatch")
            return sum(
                OpenXmlEditor._replace_in_container(box, expected, replacement)
                for box in text_boxes
            )
        raise BindingAuthorizationError("certified field has unsupported element kind")

    def _validate_all_table_locators(self) -> None:
        body_nodes = tuple(self.document.element.body.iterchildren())
        seen: set[tuple[Any, ...]] = set()
        for table_id, binding in self._tables.items():
            locator = binding["locator"]
            identity = (
                locator["story"], locator["body_node_index"], locator["table_index"],
                locator["table_signature"], binding["semantic_parent"],
            )
            if identity in seen:
                raise TemplateIntegrityError("certified table identities collide")
            seen.add(identity)
            try:
                table = self.document.tables[locator["table_index"]]._tbl
                body_node = body_nodes[locator["body_node_index"]]
            except IndexError as error:
                raise LocatorMismatch(f"table locator out of range: {table_id}") from error
            if table is not body_node:
                raise LocatorMismatch(f"body/table coordinates disagree: {table_id}")
            if _digest(table) != locator["table_signature"]:
                raise LocatorMismatch(f"table structural signature mismatch: {table_id}")

    def resolve_dynamic_cell(self, binding_id: str):
        if binding_id in self._static_cells:
            raise BindingAuthorizationError("static table cells are immutable")
        try:
            binding = self._dynamic_cells[binding_id]
        except KeyError as error:
            raise BindingAuthorizationError("cell is not a certified dynamic binding") from error
        locator = binding["locator"]
        table = self.document.tables[locator["table_index"]]
        try:
            row = table._tbl.findall(qn("w:tr"))[locator["row_index"]]
            cell = row.findall(qn("w:tc"))[locator["physical_cell_index"]]
        except IndexError as error:
            raise LocatorMismatch("dynamic cell coordinates no longer match") from error
        if _digest(cell.find(qn("w:tcPr"))) != locator["tcPr_sha256"]:
            raise LocatorMismatch("dynamic cell formatting signature mismatch")
        return cell

    def replace_dynamic_cell(self, binding_id: str, replacement: str) -> None:
        self._validate_plain_text(replacement)
        self.contracts.validate_structural_operation(
            StructuralOperation.REPLACE_CONTENT, target_id=binding_id
        )
        cell = self.resolve_dynamic_cell(binding_id)
        text_nodes = cell.xpath(".//w:t[not(ancestor::w:instrText)]")
        if text_nodes:
            text_nodes[0].text = replacement
            for node in text_nodes[1:]:
                node.text = ""
            return
        paragraph = cell.find(qn("w:p"))
        if paragraph is None:
            paragraph = OxmlElement("w:p")
            cell.append(paragraph)
        run = paragraph.find(qn("w:r"))
        if run is None:
            run = OxmlElement("w:r")
            paragraph.append(run)
        value = OxmlElement("w:t")
        value.text = replacement
        run.append(value)

    def replace_certified_manual_diagram(self, location_id: str) -> None:
        """Replace one exact project visual with the certified manual placeholder.

        This method cannot replace branding, accept arbitrary text, or add image
        content. It removes only the certified drawing and its now-unused image
        relationship while retaining the paragraph and run property containers.
        """

        try:
            certification = self._manual_diagrams[location_id]
        except KeyError as error:
            raise BindingAuthorizationError(
                "visual is not a certified manual-diagram location"
            ) from error
        if certification["classification"] != "MANUAL_DIAGRAM_PLACEHOLDER":
            raise BindingAuthorizationError("visual classification is not replaceable")
        if certification["placeholder_text"] != "[ INSERT DIAGRAM ]":
            raise TemplateIntegrityError("manual-diagram placeholder contract changed")
        locator = certification["locator"]
        body_nodes = tuple(self.document.element.body.iterchildren())
        try:
            paragraph = body_nodes[locator["body_node_index"]]
        except IndexError as error:
            raise LocatorMismatch("manual-diagram body locator is out of range") from error
        if _digest(paragraph) != locator["structural_xml_signature"]:
            raise LocatorMismatch("manual-diagram structural signature mismatch")
        drawing = self._drawing_by_name(paragraph, locator["drawing_name"])
        relationship_attribute = (
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        relationship_ids = [
            node.get(relationship_attribute)
            for node in drawing.iter(
                "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            if node.get(relationship_attribute)
        ]
        if relationship_ids != [locator["relationship_id"]]:
            raise LocatorMismatch("manual-diagram relationship does not match")
        run = drawing.getparent()
        if run is None or run.tag != qn("w:r"):
            raise LocatorMismatch("manual-diagram drawing is not contained by one run")
        if any(node.text for node in run.findall(qn("w:t"))):
            raise LocatorMismatch("manual-diagram run contains unexpected text")
        run.remove(drawing)
        text = OxmlElement("w:t")
        text.text = certification["placeholder_text"]
        run.append(text)
        relationship_id = locator["relationship_id"]
        if relationship_id not in self.document.part.rels:
            raise LocatorMismatch("manual-diagram relationship is absent")
        self.document.part.drop_rel(relationship_id)

    def clone_certified_row(self, prototype_id: str):
        try:
            prototype = self._prototypes[prototype_id]
        except KeyError as error:
            raise BindingAuthorizationError("row prototype is not certified") from error
        table_binding = self._tables[prototype["table_id"]]
        table = self.document.tables[table_binding["locator"]["table_index"]]
        try:
            row = table._tbl.findall(qn("w:tr"))[prototype["row_index"]]
        except IndexError as error:
            raise LocatorMismatch("certified row prototype is missing") from error
        if _digest(row) != prototype["row_signature"]:
            raise LocatorMismatch("certified row prototype signature mismatch")
        relationship_attributes = {
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id",
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed",
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link",
        }
        if any(set(element.attrib) & relationship_attributes for element in row.iter()):
            raise BindingAuthorizationError("relationship-bound rows cannot be cloned")
        return deepcopy(row)

    def insert_certified_row(
        self,
        prototype_id: str,
        *,
        insert_after_row_index: int,
    ) -> CertifiedClonedRow:
        """Insert a certified clone inside its owning table and nowhere else."""

        clone = self.clone_certified_row(prototype_id)
        prototype = self._prototypes[prototype_id]
        table_binding = self._tables[prototype["table_id"]]
        table = self.document.tables[table_binding["locator"]["table_index"]]._tbl
        rows = table.findall(qn("w:tr"))
        if insert_after_row_index < 0 or insert_after_row_index >= len(rows):
            raise LocatorMismatch("certified row insertion index is out of range")
        if insert_after_row_index in table_binding["header_row_indexes"]:
            raise BindingAuthorizationError("cannot insert a data row after a header boundary")
        rows[insert_after_row_index].addnext(clone)
        return CertifiedClonedRow(
            prototype_id=prototype_id,
            table_id=prototype["table_id"],
            element=clone,
        )

    def populate_certified_cloned_row(
        self,
        cloned_row: CertifiedClonedRow,
        values: tuple[str, ...],
    ) -> None:
        """Populate only the cells materialized by a certified row prototype."""

        try:
            prototype = self._prototypes[cloned_row.prototype_id]
            table = self._tables[cloned_row.table_id]
        except KeyError as error:
            raise BindingAuthorizationError("cloned row provenance is not certified") from error
        if prototype["table_id"] != cloned_row.table_id:
            raise BindingAuthorizationError("cloned row/table provenance mismatch")
        cells = cloned_row.element.findall(qn("w:tc"))
        if len(values) != table["column_count"] or len(cells) != len(values):
            raise LocatorMismatch("cloned row values do not match certified columns")
        for value in values:
            self._validate_plain_text(value)
        for cell, value in zip(cells, values):
            text_nodes = cell.xpath(".//w:t[not(ancestor::w:instrText)]")
            if not text_nodes:
                paragraph = cell.find(qn("w:p"))
                if paragraph is None:
                    raise LocatorMismatch("certified prototype cell has no paragraph container")
                run = paragraph.find(qn("w:r"))
                if run is None:
                    run = OxmlElement("w:r")
                    paragraph.append(run)
                text_node = OxmlElement("w:t")
                run.append(text_node)
                text_nodes = [text_node]
            text_nodes[0].text = value
            for node in text_nodes[1:]:
                node.text = ""

    def save(self, output_path: str | Path | None = None) -> Path:
        """Save only to a non-source DOCX and re-assert source immutability."""

        destination = Path(output_path or self.working_copy).resolve()
        if destination == self.source_template:
            raise BindingAuthorizationError("source template cannot be an output target")
        if destination.suffix.lower() != ".docx":
            raise ValueError("certified output must be a .docx")
        destination.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(destination)
        self._restore_immutable_package_parts(destination)
        if file_sha256(self.source_template) != self.source_hash:
            raise TemplateIntegrityError("source template changed during certified population")
        return destination

    def _restore_immutable_package_parts(self, destination: Path) -> None:
        """Undo serializer-only churn outside certified mutable OPC stories."""

        mutable_parts = {
            "word/document.xml",
            "word/_rels/document.xml.rels",
            *(f"word/header{index}.xml" for index in range(1, 6)),
        }
        handle = tempfile.NamedTemporaryFile(
            prefix="phase4c-certified-", suffix=".docx", dir=destination.parent, delete=False
        )
        temporary = Path(handle.name)
        handle.close()
        try:
            with ZipFile(self.source_template) as source, ZipFile(destination) as populated:
                source_names = set(source.namelist())
                with ZipFile(temporary, "w", compression=ZIP_DEFLATED) as hardened:
                    for output_info in populated.infolist():
                        name = output_info.filename
                        payload = populated.read(name)
                        if name in source_names and name not in mutable_parts:
                            payload = source.read(name)
                        base_info = source.getinfo(name) if name in source_names else output_info
                        info = ZipInfo(name, date_time=base_info.date_time)
                        info.compress_type = base_info.compress_type
                        info.comment = base_info.comment
                        info.extra = base_info.extra
                        info.internal_attr = base_info.internal_attr
                        info.external_attr = base_info.external_attr
                        info.create_system = base_info.create_system
                        hardened.writestr(info, payload)
            os.replace(temporary, destination)
        finally:
            if temporary.exists():
                temporary.unlink()


__all__ = [
    "BindingAuthorizationError",
    "CertifiedClonedRow",
    "CertifiedContracts",
    "CertifiedTemplateBinding",
    "StructuralOperation",
]

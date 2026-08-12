"""Narrow, fail-closed editing primitives for template-preserving DOCX work."""

from __future__ import annotations

from copy import deepcopy
from hashlib import sha256
from pathlib import Path
import re
import shutil
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .template_map import ElementKind, Story, StructuralLocator, TemplateSemanticMap

MC_ALTERNATE_CONTENT = "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent"
WP_DOCPR = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr"
WPS_CNVPR = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}cNvPr"
VML_SHAPE = "{urn:schemas-microsoft-com:vml}shape"
RELATIONSHIP_ATTRIBUTES = {
    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed",
    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link",
    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id",
}


class TemplateIntegrityError(ValueError):
    """The supplied package does not match its versioned semantic map."""


class LocatorMismatch(ValueError):
    """A structural locator did not resolve exactly as declared."""


class UnsafeOpenXmlOperation(ValueError):
    """An edit could damage relationships, fields, or section structure."""


def file_sha256(path: str | Path) -> str:
    return sha256(Path(path).read_bytes()).hexdigest()


def copy_template(source: str | Path, working_copy: str | Path) -> Path:
    """Create a distinct working DOCX without ever opening the source for writing."""

    source_path = Path(source).resolve()
    target = Path(working_copy).resolve()
    if source_path == target:
        raise ValueError("working copy must differ from source template")
    if source_path.suffix.lower() != ".docx" or target.suffix.lower() != ".docx":
        raise ValueError("source and working copy must be .docx files")
    if not source_path.is_file():
        raise FileNotFoundError(source_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_path, target)
    return target


class OpenXmlEditor:
    """Edit a copied DOCX using python-docx plus isolated OpenXML operations."""

    def __init__(
        self,
        source_template: str | Path,
        working_copy: str | Path,
        semantic_map: TemplateSemanticMap,
    ) -> None:
        self.source_template = Path(source_template).resolve()
        self.source_hash = file_sha256(self.source_template)
        if self.source_hash != semantic_map.template_sha256:
            raise TemplateIntegrityError("source template hash does not match semantic map")
        self.working_copy = copy_template(self.source_template, working_copy)
        self.semantic_map = semantic_map
        self.document: DocumentObject = Document(self.working_copy)

    def save(self, output_path: str | Path) -> Path:
        destination = Path(output_path).resolve()
        if destination == self.source_template:
            raise UnsafeOpenXmlOperation("source template cannot be an output target")
        if destination.suffix.lower() != ".docx":
            raise ValueError("output path must have a .docx suffix")
        destination.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(destination)
        self.assert_source_unchanged()
        return destination

    def assert_source_unchanged(self) -> None:
        if file_sha256(self.source_template) != self.source_hash:
            raise TemplateIntegrityError("source template changed during editing")

    @property
    def body_nodes(self) -> tuple[object, ...]:
        return tuple(self.document.element.body.iterchildren())

    def resolve_body_node(self, locator: StructuralLocator):
        if locator.story is not Story.BODY or locator.body_node_index is None:
            raise LocatorMismatch("locator does not identify a body node")
        nodes = self.body_nodes
        try:
            node = nodes[locator.body_node_index]
        except IndexError as error:
            raise LocatorMismatch("body node index is out of range") from error
        self._assert_expected_text(node, locator.expected_text)
        return node

    def resolve_table_cell(self, locator: StructuralLocator):
        if locator.element_kind is not ElementKind.TABLE_CELL:
            raise LocatorMismatch("locator is not a table cell")
        assert locator.table_index is not None and locator.row_index is not None and locator.cell_index is not None
        try:
            table = self.document.tables[locator.table_index]
            row = table.rows[locator.row_index]
            cell = row.cells[locator.cell_index]
        except IndexError as error:
            raise LocatorMismatch("table-cell coordinates are out of range") from error
        if locator.surrounding_label is not None:
            label = "".join(row.cells[0]._tc.xpath(".//w:t/text()"))
            if self._normalize(label) != self._normalize(locator.surrounding_label):
                raise LocatorMismatch("table-cell surrounding label does not match")
        if locator.expected_text is not None:
            self._assert_expected_text(cell._tc, locator.expected_text)
        return cell

    def replace_text(self, locator: StructuralLocator, replacement: str) -> int:
        if locator.element_kind is ElementKind.TABLE_CELL:
            self.replace_table_cell_text(locator, replacement)
            return 1
        if locator.element_kind is ElementKind.PARAGRAPH:
            node = self.resolve_body_node(locator)
            if locator.expected_text is None:
                raise LocatorMismatch("paragraph replacement requires expected_text")
            return self._replace_in_container(node, locator.expected_text, replacement)
        if locator.element_kind is ElementKind.DRAWING_TEXT:
            return self.replace_drawing_text(locator, replacement)
        raise LocatorMismatch("locator does not address replaceable text")

    def replace_table_cell_text(self, locator: StructuralLocator, replacement: str) -> None:
        cell = self.resolve_table_cell(locator)
        paragraphs = cell._tc.xpath("./w:p")
        if not paragraphs:
            paragraph = OxmlElement("w:p")
            cell._tc.append(paragraph)
            paragraphs = [paragraph]
        text_nodes = self._editable_text_nodes(cell._tc)
        current = "".join(node.text or "" for node in text_nodes)
        if locator.expected_text is not None:
            if self._normalize(current) != self._normalize(locator.expected_text):
                raise LocatorMismatch("table-cell value does not match expected text")
            self._replace_across_nodes(text_nodes, current, replacement)
        elif current:
            self._replace_across_nodes(text_nodes, current, replacement)
        else:
            run = paragraphs[0].find(qn("w:r"))
            if run is None:
                run = OxmlElement("w:r")
                paragraphs[0].append(run)
            text = OxmlElement("w:t")
            text.text = replacement
            run.append(text)

    def clone_table_row(self, table_index: int, row_index: int, insert_after: int | None = None):
        try:
            table = self.document.tables[table_index]
            source = table.rows[row_index]._tr
        except IndexError as error:
            raise LocatorMismatch("table row coordinates are out of range") from error
        if self._has_relationship_attributes(source):
            raise UnsafeOpenXmlOperation("relationship-bound rows require explicit remapping")
        clone = deepcopy(source)
        reference = table.rows[row_index if insert_after is None else insert_after]._tr
        reference.addnext(clone)
        return clone

    def replace_drawing_text(self, locator: StructuralLocator, replacement: str) -> int:
        root = self._story_root(locator)
        drawing = self._drawing_by_name(root, locator.drawing_name or "")
        scope = drawing
        for ancestor in drawing.iterancestors():
            if ancestor.tag == MC_ALTERNATE_CONTENT:
                scope = ancestor
                break
        text_boxes = list(scope.iter(qn("w:txbxContent")))
        if not text_boxes:
            raise LocatorMismatch("drawing has no text-box content")
        expected = locator.expected_text
        if expected is None:
            raise LocatorMismatch("drawing replacement requires expected_text")
        if any(expected not in self._all_text(text_box) for text_box in text_boxes):
            raise LocatorMismatch("DrawingML and VML text payloads were not both matched")
        return sum(
            self._replace_in_container(text_box, expected, replacement)
            for text_box in text_boxes
        )

    def clone_banner(self, locator: StructuralLocator):
        """Clone a relationship-free banner paragraph and remap drawing/VML IDs."""

        node = self.resolve_body_node(locator)
        drawing = self._drawing_by_name(node, locator.drawing_name or "")
        if self._has_relationship_attributes(node):
            raise UnsafeOpenXmlOperation("relationship-bound drawings require explicit remapping")
        clone = deepcopy(node)
        self._remap_drawing_ids(clone)
        return clone

    def clone_body_node(self, locator: StructuralLocator):
        """Deep-clone a verified relationship-free paragraph, table, or drawing node."""

        node = self.resolve_body_node(locator)
        if locator.table_index is not None:
            try:
                expected_table = self.document.tables[locator.table_index]._tbl
            except IndexError as error:
                raise LocatorMismatch("prototype table index is out of range") from error
            if node is not expected_table:
                raise LocatorMismatch("body node and table prototype indexes disagree")
        if self._has_relationship_attributes(node):
            raise UnsafeOpenXmlOperation("relationship-bound prototypes require explicit remapping")
        clone = deepcopy(node)
        if any(element.tag == WP_DOCPR for element in clone.iter()):
            self._remap_drawing_ids(clone)
        return clone

    def clone_detached_node(self, node):
        """Clone a captured prototype and allocate IDs against the current document."""

        if node.getparent() is not None:
            raise UnsafeOpenXmlOperation("prototype clone source must be detached")
        if self._has_relationship_attributes(node):
            raise UnsafeOpenXmlOperation("relationship-bound prototypes require explicit remapping")
        clone = deepcopy(node)
        if any(element.tag == WP_DOCPR for element in clone.iter()):
            self._remap_drawing_ids(clone)
        return clone

    def replace_text_in_node(self, node, expected: str, replacement: str) -> int:
        """Run-aware replacement inside an already cloned OpenXML node."""

        return self._replace_in_container(node, expected, replacement)

    def replace_all_text_in_node(self, node, replacement: str) -> None:
        """Replace all non-field text while retaining the existing run/property tree."""

        text_nodes = self._editable_text_nodes(node)
        current = "".join(text.text or "" for text in text_nodes)
        if not current:
            raise LocatorMismatch("node contains no editable text")
        self._replace_across_nodes(text_nodes, current, replacement)

    def replace_cloned_drawing_text(
        self, node, drawing_name: str, expected: str | None, replacement: str,
    ) -> int:
        """Synchronize text across DrawingML and VML branches in a detached clone."""

        drawing = self._drawing_by_name(node, drawing_name)
        scope = drawing
        for ancestor in drawing.iterancestors():
            if ancestor.tag == MC_ALTERNATE_CONTENT:
                scope = ancestor
                break
        text_boxes = list(scope.iter(qn("w:txbxContent")))
        if not text_boxes:
            raise LocatorMismatch("drawing has no text-box content")
        if expected is not None:
            if any(expected not in self._all_text(text_box) for text_box in text_boxes):
                raise LocatorMismatch("cloned DrawingML and VML payloads do not match")
            return sum(self._replace_in_container(box, expected, replacement) for box in text_boxes)
        for text_box in text_boxes:
            self.replace_all_text_in_node(text_box, replacement)
        return len(text_boxes)

    def populate_table(
        self,
        table_node,
        headers: tuple[str, ...],
        rows: tuple[tuple[str, ...], ...],
        *,
        header_row_index: int,
        data_row_index: int,
        alternate_data_row_index: int | None = None,
    ) -> None:
        """Populate a cloned table using its existing header and data-row prototypes."""

        table_rows = list(table_node.findall(qn("w:tr")))
        prototype_indexes = (data_row_index,) + (
            (alternate_data_row_index,) if alternate_data_row_index is not None else ()
        )
        if not table_rows or max((header_row_index, *prototype_indexes)) >= len(table_rows):
            raise LocatorMismatch("table prototype row indexes are invalid")
        header = deepcopy(table_rows[header_row_index])
        data_prototypes = tuple(deepcopy(table_rows[index]) for index in prototype_indexes)
        if len(header.findall(qn("w:tc"))) != len(headers):
            raise LocatorMismatch("semantic headers do not match prototype columns")
        if any(len(data.findall(qn("w:tc"))) != len(headers) for data in data_prototypes):
            raise LocatorMismatch("semantic rows do not match prototype columns")
        for existing in table_rows:
            table_node.remove(existing)
        table_node.append(header)
        self._populate_row(header, headers)
        for row_index, values in enumerate(rows):
            if len(values) != len(headers):
                raise LocatorMismatch("semantic table row width changed during publishing")
            row = deepcopy(data_prototypes[row_index % len(data_prototypes)])
            self._populate_row(row, values)
            table_node.append(row)

    def apply_existing_numbering(
        self, paragraph_node, *, num_id: int = 1, level: int = 3, expected_format: str = "bullet",
    ) -> None:
        """Apply an existing numbering definition after verifying its level format."""

        numbering = self.document.part.numbering_part.element
        nums = numbering.xpath(f'./w:num[@w:numId="{num_id}"]/w:abstractNumId/@w:val')
        if len(nums) != 1:
            raise LocatorMismatch("numbering instance is not present in the template")
        formats = numbering.xpath(
            f'./w:abstractNum[@w:abstractNumId="{nums[0]}"]/w:lvl[@w:ilvl="{level}"]/w:numFmt/@w:val'
        )
        if formats != [expected_format]:
            raise LocatorMismatch("numbering level does not provide the required format")
        p_pr = paragraph_node.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            paragraph_node.insert(0, p_pr)
        old = p_pr.find(qn("w:numPr"))
        if old is not None:
            p_pr.remove(old)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        number = OxmlElement("w:numId")
        number.set(qn("w:val"), str(num_id))
        num_pr.extend((ilvl, number))
        p_pr.append(num_pr)

    def clear_dynamic_body(self, start_body_node: int) -> tuple[object, ...]:
        """Remove project content after the cover while retaining every section boundary."""

        nodes = self.body_nodes
        if start_body_node < 0 or start_body_node >= len(nodes):
            raise LocatorMismatch("dynamic body start is out of range")
        body = self.document.element.body
        preserved: list[object] = []
        for node in nodes[start_body_node:]:
            if node.tag == qn("w:sectPr") or node.xpath("./w:pPr/w:sectPr"):
                preserved.append(node)
            else:
                body.remove(node)
        if not preserved or preserved[-1].tag != qn("w:sectPr"):
            raise TemplateIntegrityError("final section properties were not preserved")
        return tuple(preserved)

    def insert_body_node(self, reference, node, position: str):
        body = self.document.element.body
        if reference.getparent() is not body:
            raise UnsafeOpenXmlOperation("reference must be a direct body node")
        if node.getparent() is not None:
            raise UnsafeOpenXmlOperation("inserted node must be a detached clone")
        if node.tag == qn("w:sectPr") or node.xpath("./w:pPr/w:sectPr"):
            raise UnsafeOpenXmlOperation("section-property nodes require a section-aware operation")
        if reference.tag == qn("w:sectPr") and position == "after":
            raise UnsafeOpenXmlOperation("no body node may follow the final sectPr")
        if position == "before":
            reference.addprevious(node)
        elif position == "after":
            reference.addnext(node)
        else:
            raise ValueError("position must be 'before' or 'after'")
        return node

    def remove_body_node(self, node) -> None:
        body = self.document.element.body
        if node.getparent() is not body:
            raise UnsafeOpenXmlOperation("removed node must be a direct body node")
        if node.tag == qn("w:sectPr") or node.xpath("./w:pPr/w:sectPr"):
            raise UnsafeOpenXmlOperation("section-property nodes cannot be removed")
        if self._has_relationship_attributes(node):
            raise UnsafeOpenXmlOperation("relationship-bound nodes require relationship cleanup")
        node.getparent().remove(node)

    def assert_footers_preserved(self, branding_text: str) -> None:
        assert_footer_contract(self.document, branding_text)

    @staticmethod
    def _has_relationship_attributes(node) -> bool:
        return any(name in RELATIONSHIP_ATTRIBUTES for element in node.iter() for name in element.attrib)

    def _story_root(self, locator: StructuralLocator):
        if locator.story is Story.BODY:
            return self.resolve_body_node(locator)
        expected = "/" + (locator.part_name or "").lstrip("/")
        matches = tuple(part for part in self.document.part.package.parts if str(part.partname) == expected)
        if len(matches) != 1:
            raise LocatorMismatch(f"story part did not resolve uniquely: {locator.part_name}")
        return matches[0].element

    @staticmethod
    def _drawing_by_name(root, name: str):
        matches = root.xpath(f'.//w:drawing[.//wp:docPr[@name="{name}"]]')
        if len(matches) != 1:
            raise LocatorMismatch(f"drawing did not resolve uniquely: {name}")
        return matches[0]

    @classmethod
    def _replace_in_container(cls, container, expected: str, replacement: str) -> int:
        nodes = cls._editable_text_nodes(container)
        combined = "".join(node.text or "" for node in nodes)
        occurrences = combined.count(expected)
        if occurrences != 1:
            raise LocatorMismatch(f"expected text occurrence count was {occurrences}, not 1")
        cls._replace_across_nodes(nodes, expected, replacement)
        return 1

    @staticmethod
    def _editable_text_nodes(container) -> list:
        # w:t never contains field instructions, but exclude cached field results as well.
        result = []
        field_depth = 0
        for element in container.iter():
            if element.tag == qn("w:fldChar"):
                kind = element.get(qn("w:fldCharType"))
                if kind == "begin":
                    field_depth += 1
                elif kind == "end" and field_depth:
                    field_depth -= 1
            elif element.tag == qn("w:t") and field_depth == 0:
                result.append(element)
        return result

    @staticmethod
    def _replace_across_nodes(nodes: Iterable, expected: str, replacement: str) -> None:
        nodes = list(nodes)
        combined = "".join(node.text or "" for node in nodes)
        start = combined.find(expected)
        if start < 0:
            raise LocatorMismatch("expected text was not found")
        end = start + len(expected)
        positions: list[tuple[object, int, int]] = []
        cursor = 0
        for node in nodes:
            next_cursor = cursor + len(node.text or "")
            if next_cursor > start and cursor < end:
                positions.append((node, cursor, next_cursor))
            cursor = next_cursor
        if not positions:
            raise LocatorMismatch("expected text has no editable run nodes")
        weights = [min(stop, end) - max(begin, start) for _, begin, stop in positions]
        total = sum(weights)
        consumed = 0
        replacement_offset = 0
        for index, ((node, begin, stop), weight) in enumerate(zip(positions, weights)):
            old = node.text or ""
            left = old[: max(0, start - begin)] if index == 0 else ""
            right = old[min(len(old), end - begin) :] if index == len(positions) - 1 else ""
            consumed += weight
            next_offset = round(len(replacement) * consumed / total)
            node.text = left + replacement[replacement_offset:next_offset] + right
            if node.text.startswith((" ", "\t", "\n")) or node.text.endswith((" ", "\t", "\n")):
                node.set(qn("xml:space"), "preserve")
            replacement_offset = next_offset

    def _remap_drawing_ids(self, clone) -> None:
        root = self.document.element
        drawing_id_tags = {WP_DOCPR, WPS_CNVPR}
        numeric = [
            int(element.get("id")) for element in root.iter()
            if element.tag in drawing_id_tags and str(element.get("id", "")).isdigit()
        ]
        next_id = max(numeric, default=0) + 1
        mapping: dict[str, str] = {}
        for element in clone.iter():
            if element.tag not in drawing_id_tags:
                continue
            old = element.get("id")
            if old not in mapping:
                mapping[old] = str(next_id)
                next_id += 1
            element.set("id", mapping[old])
        vml_numbers = [
            int(match.group(1)) for element in root.iter()
            if element.tag == VML_SHAPE
            if (match := re.fullmatch(r"docshape(\d+)", element.get("id", "")))
        ]
        next_vml = max(vml_numbers, default=0) + 1
        for shape in clone.iter():
            if shape.tag == VML_SHAPE:
                shape.set("id", f"docshape{next_vml}")
                next_vml += 1

    def _populate_row(self, row, values: tuple[str, ...]) -> None:
        cells = row.findall(qn("w:tc"))
        if len(cells) != len(values):
            raise LocatorMismatch("row prototype and semantic values have different widths")
        for cell, value in zip(cells, values):
            text_nodes = self._editable_text_nodes(cell)
            current = "".join(text.text or "" for text in text_nodes)
            if current:
                self._replace_across_nodes(text_nodes, current, value)
                continue
            paragraphs = cell.findall(qn("w:p"))
            paragraph = paragraphs[0] if paragraphs else OxmlElement("w:p")
            if not paragraphs:
                cell.append(paragraph)
            run = paragraph.find(qn("w:r"))
            if run is None:
                run = OxmlElement("w:r")
                paragraph.append(run)
            text = OxmlElement("w:t")
            text.text = value
            run.append(text)

    @staticmethod
    def _assert_expected_text(container, expected: str | None) -> None:
        if expected is None:
            return
        actual = OpenXmlEditor._all_text(container)
        if OpenXmlEditor._normalize(expected) not in OpenXmlEditor._normalize(actual):
            raise LocatorMismatch("resolved element does not contain expected text")

    @staticmethod
    def _normalize(value: str) -> str:
        return " ".join(value.split())

    @staticmethod
    def _all_text(container) -> str:
        return "".join(element.text or "" for element in container.iter(qn("w:t")))


def assert_footer_contract(document: DocumentObject, branding_text: str) -> None:
    """Assert that every audited footer retains drawings, branding, and page fields."""

    footers = tuple(
        part for part in document.part.package.parts
        if str(part.partname).startswith("/word/footer")
    )
    if len(footers) != 5:
        raise TemplateIntegrityError("expected five footer parts")
    for part in footers:
        instructions = {text.strip() for text in part.element.xpath(".//w:instrText/text()")}
        if not {"PAGE", "NUMPAGES"}.issubset(instructions):
            raise TemplateIntegrityError(f"footer fields missing from {part.partname}")
        if len(part.element.xpath(".//w:drawing")) != 2:
            raise TemplateIntegrityError(f"footer drawings changed in {part.partname}")
        if branding_text not in "".join(part.element.xpath(".//w:t/text()")):
            raise TemplateIntegrityError(f"footer branding missing from {part.partname}")


__all__ = [
    "LocatorMismatch", "OpenXmlEditor", "TemplateIntegrityError", "assert_footer_contract",
    "UnsafeOpenXmlOperation", "copy_template", "file_sha256",
]

"""Deterministic DOCX rendering for immutable proposal-domain documents."""

from __future__ import annotations

import re
from collections.abc import Mapping
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.document import Document

from .contracts import ProposalDocument


class DOCXRenderer:
    """Render proposal documents using Word's built-in document styles."""

    def render(
        self,
        proposal_document: ProposalDocument,
        output_path: str | Path,
        *,
        template_path: str | Path | None = None,
    ) -> Path:
        """Write one professional DOCX without changing proposal-domain content."""
        if not isinstance(proposal_document, ProposalDocument):
            raise TypeError("proposal_document must be a ProposalDocument")
        destination = Path(output_path)
        if destination.suffix.lower() != ".docx":
            raise ValueError("output_path must have a .docx suffix")
        document = self._load_document(template_path)
        self._apply_template_placeholders(document, {"proposal_title": proposal_document.proposal_title})
        self._apply_header_footer(document, proposal_document.metadata)

        if not self._contains_text(document, proposal_document.proposal_title):
            document.add_heading(proposal_document.proposal_title, level=0)

        plans_by_id = {section.section_id: section for section in proposal_document.proposal_plan.sections}
        for index, section_content in enumerate(proposal_document.sections):
            section_plan = plans_by_id[section_content.section_id]
            if index and bool(section_plan.metadata.get("page_break_before", False)):
                document.add_page_break()
            document.add_heading(section_plan.title, level=1)
            self._render_content(document, section_content.generated_text)

        if proposal_document.references:
            document.add_heading("References", level=1)
            for reference in proposal_document.references:
                document.add_paragraph(
                    f"[{reference.reference_id}] {reference.source_document} — {reference.content}",
                    style="List Bullet",
                )

        destination.parent.mkdir(parents=True, exist_ok=True)
        document.save(destination)
        return destination

    @staticmethod
    def _load_document(template_path: str | Path | None) -> Document:
        if template_path is None:
            return DocxDocument()
        template = Path(template_path)
        if template.suffix.lower() != ".docx":
            raise ValueError("template_path must have a .docx suffix")
        if not template.is_file():
            raise FileNotFoundError(template)
        return DocxDocument(template)

    @staticmethod
    def _apply_template_placeholders(document: Document, values: Mapping[str, str]) -> None:
        placeholders = {f"{{{{{name}}}}}": value for name, value in values.items()}
        for paragraph in DOCXRenderer._all_paragraphs(document):
            if any(token in paragraph.text for token in placeholders):
                paragraph.text = DOCXRenderer._replace_values(paragraph.text, placeholders)

    @staticmethod
    def _apply_header_footer(document: Document, metadata: Mapping[str, Any]) -> None:
        header_text = metadata.get("header_text")
        footer_text = metadata.get("footer_text")
        for section in document.sections:
            if isinstance(header_text, str):
                DOCXRenderer._set_story_text(section.header.paragraphs, header_text)
            if isinstance(footer_text, str):
                DOCXRenderer._set_story_text(section.footer.paragraphs, footer_text)

    @staticmethod
    def _set_story_text(paragraphs: list[Any], text: str) -> None:
        if paragraphs:
            paragraphs[0].text = text

    @staticmethod
    def _contains_text(document: Document, text: str) -> bool:
        return any(paragraph.text == text for paragraph in DOCXRenderer._all_paragraphs(document))

    @staticmethod
    def _all_paragraphs(document: Document):
        for paragraph in document.paragraphs:
            yield paragraph
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
        for section in document.sections:
            yield from section.header.paragraphs
            yield from section.footer.paragraphs

    @staticmethod
    def _replace_values(text: str, placeholders: Mapping[str, str]) -> str:
        for token, value in placeholders.items():
            text = text.replace(token, value)
        return text

    def _render_content(self, document: Document, content: str) -> None:
        lines = content.splitlines()
        index = 0
        while index < len(lines):
            table_end = self._markdown_table_end(lines, index)
            if table_end is not None:
                self._render_markdown_table(document, lines[index:table_end])
                index = table_end
            else:
                document.add_paragraph(lines[index])
                index += 1
        if not lines:
            document.add_paragraph("")

    @staticmethod
    def _markdown_table_end(lines: list[str], start: int) -> int | None:
        if start + 1 >= len(lines) or not DOCXRenderer._is_table_row(lines[start]):
            return None
        if not DOCXRenderer._is_separator_row(lines[start + 1]):
            return None
        column_count = len(DOCXRenderer._table_cells(lines[start]))
        end = start + 2
        while end < len(lines) and DOCXRenderer._is_table_row(lines[end]):
            if len(DOCXRenderer._table_cells(lines[end])) != column_count:
                break
            end += 1
        return end

    @staticmethod
    def _is_table_row(line: str) -> bool:
        return "|" in line and len(DOCXRenderer._table_cells(line)) >= 2

    @staticmethod
    def _is_separator_row(line: str) -> bool:
        cells = DOCXRenderer._table_cells(line)
        return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) is not None for cell in cells)

    @staticmethod
    def _table_cells(line: str) -> list[str]:
        trimmed = line.strip().strip("|")
        return [cell.strip() for cell in trimmed.split("|")]

    def _render_markdown_table(self, document: Document, lines: list[str]) -> None:
        rows = [self._table_cells(line) for position, line in enumerate(lines) if position != 1]
        table = document.add_table(rows=1, cols=len(rows[0]))
        table.style = "Light Shading Accent 1"
        for row_index, values in enumerate(rows):
            cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
            for cell, value in zip(cells, values):
                cell.text = value

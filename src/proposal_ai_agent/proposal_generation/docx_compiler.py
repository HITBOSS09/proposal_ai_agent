"""Compile semantic composition documents into template-owned DOCX files."""

from __future__ import annotations

from collections.abc import Mapping
from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document

from .composition import ComponentContent, ComponentInstance, CompositionDocument
from .proposal_ir import (
    BulletList,
    Callout,
    Heading,
    Paragraph,
    RequirementMatrix,
    Table,
    VisualPlaceholder,
)
from .publishing import PRU_TEMPLATE_SHA256, TemplatePublisher, file_sha256, pru_template_semantic_map
from .word_style_contract import WordStyleContract


class MissingStyleContract(KeyError):
    """A semantic compiler style key was not supplied."""


class MissingTemplateStyle(ValueError):
    """The master template does not define a requested named style."""


class DOCXCompiler:
    """Translate semantic component instances into a master-template DOCX."""

    def __init__(self, style_contracts: Mapping[str, WordStyleContract]) -> None:
        self._style_contracts = dict(style_contracts)

    def compile(
        self,
        composition: CompositionDocument,
        master_template_path: str | Path,
        output_path: str | Path,
    ) -> Path:
        """Compile one composition document using an existing master template."""
        if not isinstance(composition, CompositionDocument):
            raise TypeError("composition must be a CompositionDocument")
        template = self._template_path(master_template_path)
        destination = Path(output_path)
        if destination.suffix.lower() != ".docx":
            raise ValueError("output_path must have a .docx suffix")

        if file_sha256(template) == PRU_TEMPLATE_SHA256:
            return TemplatePublisher().publish(
                composition,
                pru_template_semantic_map(template),
                template,
                destination,
            )

        document = DocxDocument(template)
        self._validate_template_styles(document)
        for component in composition.components:
            self._render_component(document, component, composition)
        destination.parent.mkdir(parents=True, exist_ok=True)
        document.save(destination)
        return destination

    @staticmethod
    def _template_path(path: str | Path) -> Path:
        template = Path(path)
        if template.suffix.lower() != ".docx":
            raise ValueError("master_template_path must have a .docx suffix")
        if not template.is_file():
            raise FileNotFoundError(template)
        return template

    def _validate_template_styles(self, document: Document) -> None:
        for contract in self._style_contracts.values():
            try:
                document.styles[contract.template_style_name]
            except KeyError as error:
                raise MissingTemplateStyle(
                    f"master template does not define style: {contract.template_style_name}"
                ) from error

    def _render_component(
        self,
        document: Document,
        component: ComponentInstance,
        composition: CompositionDocument,
    ) -> None:
        if self._metadata_flag(component, "page_break_before"):
            document.add_page_break()

        if component.component_name == "cover_page":
            self._add_paragraph(document, composition.title, "proposal_title")
        for slot in component.slots:
            for content in slot.contents:
                self._render_content(document, content, component.component_name)
        if component.component_name == "references":
            for reference in composition.references:
                self._add_paragraph(
                    document,
                    " — ".join(value for value in (reference.reference_id, reference.title, reference.source, reference.locator) if value),
                    "body",
                )
        for child in component.children:
            self._render_component(document, child, composition)

    def _render_content(self, document: Document, content: ComponentContent, component_name: str) -> None:
        node = content.content
        if isinstance(node, Heading):
            self._add_paragraph(document, node.text, self._heading_style_key(node.level, component_name))
        elif isinstance(node, Paragraph):
            self._add_paragraph(document, node.text, "body")
        elif isinstance(node, BulletList):
            for item in node.items:
                self._add_paragraph(document, item, "bullet")
        elif isinstance(node, Table):
            self._render_table(document, node)
        elif isinstance(node, RequirementMatrix):
            self._render_requirement_matrix(document, node)
        elif isinstance(node, VisualPlaceholder):
            self._add_paragraph(document, node.description, "body")
            if node.caption is not None:
                self._add_paragraph(document, node.caption, "caption")
        elif isinstance(node, Callout):
            self._add_paragraph(document, node.label, "callout")
            self._add_paragraph(document, node.text, "callout")
        else:
            raise TypeError("content must contain a supported semantic node")

    def _render_table(self, document: Document, node: Table) -> None:
        table = document.add_table(rows=1, cols=len(node.headers))
        for index, header in enumerate(node.headers):
            self._set_cell_text(table.cell(0, index), header, "table_header")
        for row in node.rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                self._set_cell_text(cells[index], value, "table_cell")

    def _render_requirement_matrix(self, document: Document, node: RequirementMatrix) -> None:
        table = document.add_table(rows=0, cols=4)
        for entry in node.entries:
            cells = table.add_row().cells
            values = (
                entry.requirement_id,
                entry.requirement,
                entry.response,
                ", ".join(entry.evidence_reference_ids),
            )
            for index, value in enumerate(values):
                self._set_cell_text(cells[index], value, "requirement_matrix")

    def _set_cell_text(self, cell, text: str, style_key: str) -> None:
        paragraph = cell.paragraphs[0]
        paragraph.text = text
        paragraph.style = self._template_style_name(style_key)

    def _add_paragraph(self, document: Document, text: str, style_key: str) -> None:
        paragraph = document.add_paragraph(text)
        paragraph.style = self._template_style_name(style_key)

    def _template_style_name(self, style_key: str) -> str:
        try:
            return self._style_contracts[style_key].template_style_name
        except KeyError as error:
            raise MissingStyleContract(f"missing style contract: {style_key}") from error

    @staticmethod
    def _metadata_flag(component: ComponentInstance, key: str) -> bool:
        return bool(dict(component.metadata).get(key, False))

    @staticmethod
    def _heading_style_key(level: int, component_name: str) -> str:
        if component_name == "cover_page":
            return "cover_title"
        if component_name == "module_banner":
            return "module_banner"
        if level == 1:
            return "heading_1"
        if level == 2:
            return "heading_2"
        if level == 3:
            return "heading_3"
        raise MissingStyleContract(f"missing semantic heading style for outline level: {level}")

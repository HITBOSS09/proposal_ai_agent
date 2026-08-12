"""
Document Parser - Converts python-docx Document to internal Document model.

This module provides the core parsing logic for converting loaded Word documents
into our production-ready internal representation. The parser preserves:
- Exact reading order (order_index)
- Hierarchical section structure
- Parent references (document_id, section_id)
- Source tracking (source_file, source_path)

Algorithm Overview:
-------------------
1. Traverse paragraphs and tables in document order
2. Detect headings by style name analysis (Heading 1, Heading 2, etc.)
3. Create Section objects hierarchically based on heading levels
4. Assign non-heading paragraphs and tables to their enclosing section
5. Track order_index continuously across all elements
6. Populate metadata fields from source document and file info

Hierarchical Nesting:
- When a Heading 2 is encountered inside a Heading 1 section, it becomes a subsection
- When a Heading 1 is encountered, previous sections close and a new top-level section opens
- Root-level elements (before any heading) stay in Document.elements

Edge Cases Handled:
- Empty paragraphs are skipped (content.strip() check)
- Documents without headings create flat structure (all elements at root level)
- Multiple top-level sections are supported
- Trailing content after last heading is preserved
"""

import logging
from pathlib import Path
from typing import List, Optional, Tuple
from uuid import UUID, uuid4

from docx.document import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from .models import (
    Document,
    DocumentMetadata,
    ElementMetadata,
    Section,
    Paragraph,
    Table,
    TableRow,
    TableCell,
    ParagraphStyle,
)

logger = logging.getLogger(__name__)
logger.addHandler(logging.NullHandler())
PARSER_VERSION = "1.0"


# HEADING DETECTION


def extract_heading_level(style_name: Optional[str]) -> Optional[int]:
    """
    Extract heading level from python-docx style name.
    
    python-docx uses style names like "Heading 1", "Heading 2", etc.
    This function extracts the numeric level, or returns None if not a heading.
    
    Args:
        style_name: The style name from a paragraph (e.g., "Heading 1")
        
    Returns:
        Integer heading level (1-6) or None if not a heading style
    """
    if not style_name:
        return None
    
    style_lower = style_name.lower().strip()
    
    # Check for standard Heading styles
    if style_lower.startswith("heading "):
        try:
            level_str = style_lower.replace("heading ", "").strip()
            level = int(level_str)
            # Validate reasonable heading range
            if 1 <= level <= 6:
                return level
        except (ValueError, IndexError):
            pass
    
    return None


def is_heading_paragraph(paragraph: DocxParagraph) -> bool:
    """
    Determine if a python-docx paragraph is a heading.
    
    Args:
        paragraph: A python-docx Paragraph object
        
    Returns:
        True if paragraph is a heading, False otherwise
    """
    return extract_heading_level(paragraph.style.name) is not None



# TABLE PARSING


def parse_table(
    docx_table: DocxTable,
    order_index: int,
    document_id: UUID,
    section_id: Optional[UUID],
    source_file: str,
    source_path: Optional[str],
    source_document: Optional[str],
) -> Table:
    """
    Convert a python-docx Table to our internal Table model.
    
    Preserves table structure: rows → cells with positional information.
    Each row and cell maintains its index for later reconstruction.
    
    Args:
        docx_table: python-docx Table object
        order_index: Position in document traversal order
        document_id: UUID of parent document
        section_id: UUID of parent section (if any)
        source_file: Filename of source document
        source_path: Full path to source file
        source_document: Logical document name/type
        
    Returns:
        Parsed Table object
    """
    rows = []
    
    for row_idx, docx_row in enumerate(docx_table.rows):
        cells = []
        
        for col_idx, docx_cell in enumerate(docx_row.cells):
            # Extract all text from cell (may have multiple paragraphs)
            cell_content = "\n".join(
                p.text for p in docx_cell.paragraphs if p.text.strip()
            )
            
            # Determine if cell is header (heuristic: first row)
            cell_type = "header" if row_idx == 0 else "data"
            
            cell = TableCell(
                content=cell_content,
                cell_type=cell_type,
                column_index=col_idx,
            )
            cells.append(cell)
        
        row = TableRow(
            cells=cells,
            row_index=row_idx,
        )
        rows.append(row)
    
    metadata = ElementMetadata(
        id=uuid4(),
        document_id=document_id,
        order_index=order_index,
        section_id=section_id,
        parent_element_id=None,
        page_number=None,
        source_file=source_file,
        source_path=source_path,
        source_document=source_document,
    )
    
    table = Table(
        metadata=metadata,
        rows=rows,
    )
    
    return table


# PARAGRAPH PARSING

def parse_paragraph(
    docx_paragraph: DocxParagraph,
    order_index: int,
    document_id: UUID,
    section_id: Optional[UUID],
    source_file: str,
    source_path: Optional[str],
    source_document: Optional[str],
) -> Optional[Paragraph]:
    """
    Convert a python-docx Paragraph to our internal Paragraph model.
    
    Skips empty paragraphs. Preserves text content exactly as-is.
    Does NOT perform text cleaning.
    
    Args:
        docx_paragraph: python-docx Paragraph object
        order_index: Position in document traversal order
        document_id: UUID of parent document
        section_id: UUID of parent section (if any)
        source_file: Filename of source document
        source_path: Full path to source file
        source_document: Logical document name/type
        
    Returns:
        Parsed Paragraph object or None if paragraph is empty
    """
    content = docx_paragraph.text
    
    # Skip empty paragraphs
    if not content.strip():
        return None
    
    metadata = ElementMetadata(
        id=uuid4(),
        document_id=document_id,
        order_index=order_index,
        section_id=section_id,
        parent_element_id=None,
        page_number=None,
        source_file=source_file,
        source_path=source_path,
        source_document=source_document,
    )
    
    paragraph = Paragraph(
        metadata=metadata,
        content=content,
        style=ParagraphStyle.NORMAL,
    )
    
    return paragraph


# ============================================================================
# DOCUMENT PARSER
# ============================================================================

class DocumentParser:
    """
    Parses python-docx Document objects into internal Document model.
    
    The parser handles:
    - Heading hierarchy and nesting
    - Content ordering and reading order preservation
    - Source metadata propagation
    - Empty paragraph filtering
    
    Usage:
        parser = DocumentParser()
        docx_doc = load_docx_document(Path("proposal.docx"))
        internal_doc = parser.parse(docx_doc, Path("proposal.docx"), "proposal")
    """
    
    def __init__(self):
        """Initialize the parser."""
        self.order_index = 0
        self.sections_stack: List[Section] = []
    
    def _get_current_section(self) -> Optional[Section]:
        """Get the current (most recent) section. Returns None if at root level."""
        return self.sections_stack[-1] if self.sections_stack else None
    
    def _push_section(self, heading: str, heading_level: int) -> Section:
        """
        Push a new section onto the stack, handling hierarchy.
        
        Algorithm:
        1. While current section level >= new heading level: pop (close sections)
        2. Create new section with appropriate level
        3. If there's a parent, add as subsection; else add to root sections
        4. Push onto stack
        
        Args:
            heading: Section heading text
            heading_level: Heading level (1-6)
            
        Returns:
            The newly created Section
        """
        # Close deeper sections when returning to higher level
        while self.sections_stack and self.sections_stack[-1].section_level >= heading_level:
            self.sections_stack.pop()
        
        new_section = Section(
            heading=heading,
            section_level=heading_level,
        )
        
        # If there's a parent section, add as subsection
        if self.sections_stack:
            parent = self.sections_stack[-1]
            parent.subsections.append(new_section)
        
        self.sections_stack.append(new_section)
        return new_section
    
    def _add_element_to_current_section_or_root(
        self,
        element,
        root_elements: list,
    ) -> None:
        """
        Add an element to the current section or to root if no section is active.
        
        Args:
            element: The element to add (Paragraph, Table, etc.)
            root_elements: List of root-level elements
        """
        current_section = self._get_current_section()
        
        if current_section:
            element.metadata.section_id = current_section.id
            current_section.elements.append(element)
        else:
            root_elements.append(element)
    
    def parse(
        self,
        docx_document: DocxDocument,
        source_path: Path,
        source_document: Optional[str] = None,
        document_title: Optional[str] = None,
        author: Optional[str] = None,
    ) -> Document:
        """
        Parse a python-docx Document into internal Document model.
        
        Algorithm:
        1. Initialize order_index counter and sections stack
        2. Iterate through all paragraphs and tables in source
        3. For each element:
           - If heading: create/manage Section (push onto stack)
           - If paragraph: parse and add to current section/root
           - If table: parse and add to current section/root
        4. Increment order_index for each non-empty element
        5. Build final Document with all sections and root elements
        
        Args:
            docx_document: python-docx Document object to parse
            source_path: Path to the source DOCX file
            source_document: Logical document name/type (e.g., 'Proposal 2024')
            document_title: Override document title (else use first heading or "Untitled")
            author: Document author
            
        Returns:
            Fully populated internal Document object
        """
        self.order_index = 0
        self.sections_stack = []
        root_elements: List = []
        top_level_sections: List[Section] = []
        
        source_path = Path(source_path).resolve()
        source_file = source_path.name
        
        # Create document with temporary metadata
        document_id = uuid4()
        
        # Infer document title
        if document_title is None:
            # Try to find first heading
            for element in docx_document.element.body:
                if element.tag.endswith("p"):
                    para = None
                    for p in docx_document.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    if para and is_heading_paragraph(para):
                        document_title = para.text
                        break
            document_title = document_title or "Untitled Document"
        
        # Parse document body
        for element in docx_document.element.body:
            # Check if it's a paragraph
            if element.tag.endswith("p"):
                # Find matching paragraph in docx_document.paragraphs
                docx_para = None
                for p in docx_document.paragraphs:
                    if p._element == element:
                        docx_para = p
                        break
                
                if docx_para is None:
                    continue
                
                # Check if it's a heading
                if is_heading_paragraph(docx_para):
                    heading_level = extract_heading_level(docx_para.style.name)
                    heading_text = docx_para.text
                    
                    # Create new section
                    section = self._push_section(heading_text, heading_level)
                    
                    # Track top-level sections for Document
                    if heading_level == 1:
                        top_level_sections.append(section)
                else:
                    # Regular paragraph
                    paragraph = parse_paragraph(
                        docx_para,
                        self.order_index,
                        document_id,
                        self._get_current_section().id
                        if self._get_current_section()
                        else None,
                        source_file,
                        str(source_path),
                        source_document,
                    )
                    
                    if paragraph:
                        self._add_element_to_current_section_or_root(
                            paragraph, root_elements
                        )
                        self.order_index += 1
            
            # Check if it's a table
            elif element.tag.endswith("tbl"):
                # Find matching table in docx_document.tables
                docx_tbl = None
                for tbl in docx_document.tables:
                    if tbl._element == element:
                        docx_tbl = tbl
                        break
                
                if docx_tbl is None:
                    continue
                
                table = parse_table(
                    docx_tbl,
                    self.order_index,
                    document_id,
                    self._get_current_section().id
                    if self._get_current_section()
                    else None,
                    source_file,
                    str(source_path),
                    source_document,
                )
                
                self._add_element_to_current_section_or_root(table, root_elements)
                self.order_index += 1
        
        # Create metadata
        metadata = DocumentMetadata(
            source=str(source_path),
            author=author,
            document_type=source_document,
        )
        # attach parser version
        metadata.custom_metadata["parser_version"] = PARSER_VERSION
        
        # Build final Document
        document = Document(
            title=document_title,
            metadata=metadata,
            document_id=document_id,
            sections=top_level_sections,
            elements=root_elements,
        )
        
        return document


__all__ = [
    "DocumentParser",
    "extract_heading_level",
    "is_heading_paragraph",
    "parse_table",
    "parse_paragraph",
]
